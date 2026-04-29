"""用无内置标题版本替换论文中的表格图片。

Word 正文中已经保留“表 5-x ...”题注，如果图片内部再画一次标题会显得重复。
该脚本复用当前表格数据，重新生成不带内部标题的 PNG，并替换前 4 个表格图片段落。
"""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "AI伪造图像识别论文_终稿.docx"
FIGURES_DIR = ROOT / "figures"


TABLES = [
    {
        "filename": "table_5_1_fsd_results_clean.png",
        "rows": [
            ["测试类别", "评估样本数", "Accuracy", "AP"],
            ["Midjourney", "9000", "79.56%", "82.04%"],
            ["Stable Diffusion", "9000", "88.34%", "91.30%"],
            ["ADM", "9000", "75.41%", "79.34%"],
            ["BigGAN", "9000", "79.27%", "82.40%"],
            ["GLIDE", "9000", "96.67%", "99.62%"],
            ["VQDM", "9000", "75.47%", "77.15%"],
        ],
        "width": 1500,
        "col_weights": [1.9, 1.2, 1.2, 1.2],
        "word_width_cm": 14.6,
    },
    {
        "filename": "table_5_2_stay_positive_results_clean.png",
        "rows": [
            ["测试类别", "模型", "Accuracy", "AP"],
            ["Midjourney", "Corvi+", "99.60%", "99.98%"],
            ["Midjourney", "Rajan/Ours+", "99.40%", "99.94%"],
            ["Stable Diffusion", "Corvi+", "99.78%", "100%"],
            ["Stable Diffusion", "Rajan/Ours+", "99.88%", "100%"],
        ],
        "width": 1500,
        "col_weights": [1.7, 1.4, 1.2, 1.2],
        "word_width_cm": 14.6,
    },
    {
        "filename": "table_5_3_fsd_training_exploration_clean.png",
        "rows": [
            ["训练设置", "Midjourney", "Stable Diffusion", "GLIDE", "VQDM", "ADM"],
            ["从零初始化训练", "71.26% / 73.80%", "75.73% / 78.73%", "98.96% / 99.60%", "84.14% / 89.77%", "63.82% / 64.99%"],
            ["官方 checkpoint 首轮微调", "81.77% / 86.05%", "95.44% / 97.95%", "99.97% / 99.99%", "98.03% / 99.20%", "75.28% / 78.54%"],
            ["第二轮保守微调", "85.49% / 90.08%", "96.03% / 98.28%", "99.96% / 99.99%", "97.82% / 99.23%", "74.13% / 76.89%"],
        ],
        "width": 1800,
        "col_weights": [1.8, 1.35, 1.65, 1.25, 1.25, 1.25],
        "word_width_cm": 15.4,
    },
    {
        "filename": "table_5_4_joint_training_results_clean.png",
        "rows": [
            ["训练阶段", "ADM Acc/AP", "SD Acc/AP", "Midjourney Acc/AP", "LVLM F1", "结论"],
            ["第一阶段：FSD + Stay-Positive", "95.50% / 97.50%", "95.34% / 97.90%", "87.00% / 90.47%", "N/A", "闭集训练稳定，主任务指标与同协议基线基本一致"],
            ["第二阶段：FSD + LVLM 最小量化", "95.22% / 97.45%", "95.33% / 97.72%", "86.97% / 90.38%", "0.1778", "语义辅助头真实进入训练，但主任务尚未显著提升"],
        ],
        "width": 1900,
        "col_weights": [2.0, 1.2, 1.2, 1.6, 0.9, 2.8],
        "word_width_cm": 15.4,
    },
]


def load_font(size: int, bold: bool = False):
    """加载中文字体，保证表格图片中的中文清晰可读。"""

    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


def split_text(text: str, font, max_width: int) -> list[str]:
    """按像素宽度换行，兼容中文、英文和百分号等混排内容。"""

    lines: list[str] = []
    current = ""
    for char in text.strip():
        candidate = current + char
        if not current or font.getlength(candidate) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines or [""]


def draw_multiline_center(draw, box, lines, font, fill) -> None:
    """在表格单元格里居中绘制多行文本。"""

    left, top, right, bottom = box
    line_height = int(font.size * 1.25)
    y = top + max(0, (bottom - top - line_height * len(lines)) // 2)
    for line in lines:
        x = left + max(0, int((right - left - font.getlength(line)) / 2))
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def render_clean_table(config: dict) -> Path:
    """生成无内部标题的表格 PNG。"""

    rows = config["rows"]
    width = config["width"]
    margin_x = 46
    margin_y = 22
    header_font = load_font(28, bold=True)
    body_font = load_font(25)
    usable_width = width - margin_x * 2
    weights = config["col_weights"]
    col_widths = [int(usable_width * weight / sum(weights)) for weight in weights]
    col_widths[-1] += usable_width - sum(col_widths)

    wrapped_rows = []
    row_heights = []
    for row_index, row in enumerate(rows):
        font = header_font if row_index == 0 else body_font
        wrapped_cells = []
        max_lines = 1
        for col_index, value in enumerate(row):
            lines = split_text(value, font, col_widths[col_index] - 26)
            wrapped_cells.append(lines)
            max_lines = max(max_lines, len(lines))
        wrapped_rows.append(wrapped_cells)
        row_heights.append(max(64 if row_index == 0 else 60, int(font.size * 1.28) * max_lines + 22))

    height = margin_y * 2 + sum(row_heights)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    x_edges = [margin_x]
    for col_width in col_widths:
        x_edges.append(x_edges[-1] + col_width)

    border = (55, 78, 102)
    header_bg = (218, 234, 247)
    alt_bg = (247, 250, 253)
    y = margin_y
    for row_index, wrapped_cells in enumerate(wrapped_rows):
        row_height = row_heights[row_index]
        bg = header_bg if row_index == 0 else (alt_bg if row_index % 2 == 0 else (255, 255, 255))
        draw.rectangle([margin_x, y, width - margin_x, y + row_height], fill=bg)
        for col_index, lines in enumerate(wrapped_cells):
            left = x_edges[col_index]
            right = x_edges[col_index + 1]
            draw.rectangle([left, y, right, y + row_height], outline=border, width=2)
            font = header_font if row_index == 0 else body_font
            draw_multiline_center(draw, (left + 8, y + 5, right - 8, y + row_height - 5), lines, font, (25, 35, 45))
        y += row_height
    draw.rectangle([margin_x, margin_y, width - margin_x, y], outline=border, width=3)

    out = FIGURES_DIR / config["filename"]
    image.save(out, quality=95)
    return out


def paragraph_has_image(paragraph) -> bool:
    """判断段落是否包含图片绘图对象。"""

    return "<w:drawing" in paragraph._p.xml


def replace_paragraph_image(paragraph, image_path: Path, width_cm: float) -> None:
    """清空图片段落并插入新的表格图片。"""

    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def main() -> None:
    """替换前 4 个表格图片段落，保留后面的图 5-4 联合训练柱状图。"""

    image_paths = [render_clean_table(config) for config in TABLES]

    doc = Document(DOCX_PATH)
    image_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph_has_image(paragraph)]
    if len(image_paragraphs) < 4:
        raise RuntimeError(f"图片段落不足，当前仅找到 {len(image_paragraphs)} 个")

    for paragraph, image_path, config in zip(image_paragraphs[:4], image_paths, TABLES):
        replace_paragraph_image(paragraph, image_path, config["word_width_cm"])

    doc.save(DOCX_PATH)
    print("replaced_table_images=4")
    for path in image_paths:
        print(path)


if __name__ == "__main__":
    main()
