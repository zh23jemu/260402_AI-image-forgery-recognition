"""修复论文 Word 表格在渲染器中被压成竖排的问题。

部分模板中的表格缺少明确的宽度、网格列宽和布局约束，artifact-tool 渲染时可能
把列压缩到极窄，表现为文字沿左侧竖排。本脚本只调整表格布局属性，不改正文内容。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "AI伪造图像识别论文_终稿.docx"


def set_or_replace(parent, child) -> None:
    """按标签替换 XML 子节点，避免重复写入多个互相冲突的宽度属性。"""

    tag = child.tag
    for old in list(parent):
        if old.tag == tag:
            parent.remove(old)
    parent.append(child)


def set_table_width(table, total_width_twips: int, column_widths_twips: list[int]) -> None:
    """为表格写入固定总宽、列宽和网格宽度，保证不同渲染器都能正常展开。"""

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total_width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    set_or_replace(tbl_pr, tbl_w)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    set_or_replace(tbl_pr, tbl_layout)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    set_or_replace(tbl_pr, jc)

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in column_widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = column_widths_twips[min(idx, len(column_widths_twips) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_or_replace(tc_pr, tc_w)

            no_wrap = OxmlElement("w:noWrap")
            # 不对长结论列启用 noWrap，避免文字过长导致右侧溢出。
            if idx != len(row.cells) - 1:
                set_or_replace(tc_pr, no_wrap)

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                    run.font.size = Pt(9)


def main() -> None:
    """统一修复当前论文中的 4 张结果表。"""

    doc = Document(DOCX_PATH)

    table_widths = [
        [3300, 2800, 2800, 2800],  # 表 5-1
        [3200, 3200, 2800, 2800],  # 表 5-2
        [3000, 2400, 2900, 2400, 2400, 2400],  # 表 5-3
        [3600, 2300, 2300, 2900, 1600, 4700],  # 表 5-4
    ]

    for table, widths in zip(doc.tables, table_widths):
        set_table_width(table, sum(widths), widths)

    doc.save(DOCX_PATH)
    print(f"fixed_tables={min(len(doc.tables), len(table_widths))}")


if __name__ == "__main__":
    main()
