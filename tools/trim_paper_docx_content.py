#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
在尽量不破坏现有版式的前提下，压缩论文正文内容。

说明：
1. 该脚本只修改指定段落的文字内容，不调整段落样式、页边距、表格和图片。
2. 为避免误覆盖，默认先在同目录生成一个备份文件。
"""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document


REPLACEMENTS: dict[int, str] = {
    15: (
        "随着扩散模型和多模态生成技术的发展，AI伪造图像检测面临更强的跨模型泛化与判别依据稳定性挑战。"
        "本文基于公开代码、预训练权重和 GenImage 数据，对 FSD、Stay-Positive 与 LVLM 三条路线开展基线复现与比较。"
        "结果表明，FSD 在 Midjourney、Stable Diffusion、ADM、BigGAN、GLIDE 和 VQDM 上分别获得 "
        "79.56%、88.34%、75.41%、79.27%、96.67% 和 75.47% 的准确率；Stay-Positive 在固定真假对照任务中表现更强，"
        "但 Rajan/Ours+ 在 ADM 上默认阈值仅得 52.25%，经 0.388818 校准后可提升至 80.77%。"
        "进一步统计发现，ADM 中存在 1137 个“Stay-Positive 判 fake 而 FSD 三组结果均判 real”的主导互补样本。"
        "综合来看，Stay-Positive 更适合作为高性能判别基线，FSD 更适合作为未知生成器泛化研究主干，LVLM 更适合作为复杂样本的语义分析与解释接口。"
    ),
    19: (
        "AI-generated images are increasingly difficult to detect across unseen generators. "
        "This thesis reproduces and compares FSD, Stay-Positive, and LVLM on GenImage using public code and pretrained checkpoints. "
        "FSD reaches 79.56%-96.67% accuracy across six generators. "
        "Stay-Positive performs better on fixed real-fake tasks, but its ADM accuracy rises from 52.25% to 80.77% after recalibration. "
        "We also observe 1137 dominant complementary ADM cases where calibrated Stay-Positive predicts fake while all FSD variants predict real. "
        "These results suggest that Stay-Positive is a strong discrimination baseline, FSD is more suitable for unknown-generator generalization, "
        "and LVLM is valuable for semantic explanation of difficult cases."
    ),
    84: (
        "近年来，GAN、扩散模型与多模态生成系统持续提升了伪造图像的真实感，图像检测任务已由寻找明显伪迹，"
        "转向应对高仿真、跨模型和跨分布场景。随着商业生成工具普及，伪造图像在内容审核、数字取证和信息安全中的风险进一步上升。"
    ),
    85: "与早期 GAN 场景不同，高质量扩散图像在全局结构与局部细节上更加自然，传统依赖纹理异常的检测方法因而明显退化。",
    86: "当前任务的核心难点主要体现在两方面：一是未知生成器条件下的泛化能力不足，二是模型可能错误依赖“真实图像特征”而非稳定伪造证据。",
    87: "因此，围绕未知生成器泛化、伪相关抑制和复杂场景解释开展系统比较，具有现实意义和研究价值。",
    89: "从理论上看，AI 伪造图像识别不仅是二分类问题，还涉及域泛化、少样本适应、伪相关抑制与可解释分析。统一复现多条技术路线，有助于判断模型究竟学到了什么。",
    90: "从应用上看，稳定的检测能力可服务于媒体审核、平台治理、新闻核验和数字取证，因此需要兼顾精度、泛化与鲁棒性。",
    92: "本文的目标不是直接提出全新模型，而是基于公开资源完成 FSD、Stay-Positive 和 LVLM 的代表性基线复现，明确不同方法在相近条件下的能力边界，并据此形成后续改进方向。",
    94: "围绕这一目标，本文主要完成五项工作：梳理相关研究；搭建 GenImage 最小可运行实验环境；复现 FSD 与 Stay-Positive 结果；整理 ADM 场景下的校准与冲突现象；讨论 LVLM 在复杂样本分析中的作用。",
    95: "研究方法上，本文采用文献分析、代码复现、结果比较和方案设计相结合的方式推进。",
    97: "本文的特点在于：先做实基线，再讨论改进；以公开代码、公开权重和公开数据作为实验基础；在结果解释上保持保守，不夸大未完成的训练与 LVLM 定量能力。",
    99: "全文依次讨论研究背景与意义、相关工作、方法定位、实验设计、结果分析、改进思路以及结论。",
    102: "AI 伪造图像检测已从早期显式伪迹分析，逐步发展到深度表征学习、泛化检测和鲁棒判别并行推进的阶段。",
    103: "当前研究重点已从“在已知数据上区分真假”转向“面对未知生成器、复杂后处理和分布变化时是否仍能稳定识别”。",
    104: "与此同时，研究者开始反思检测器的判别依据是否可靠，即模型究竟在学习伪造机制，还是在利用数据中的偶然偏差。",
    105: "因此，泛化能力、鲁棒性和可解释性成为当前研究的核心评价维度。",
    107: "面向固定生成器的检测方法通常把任务建模为标准二分类问题，在训练和测试分布相近时往往能够取得较高指标。",
    108: "这类方法的局限在于高度依赖训练分布，一旦生成器类型、分辨率或后处理方式发生变化，性能可能快速下降。",
    109: "以 CNN-Generated Images Are Surprisingly Easy to Spot... for Now 为代表的工作说明，闭集高精度并不等同于开放场景下的长期有效性。",
    111: "面向未知生成器的泛化检测方法更强调模型在新类别和新分布下的适应能力，通常结合元学习、少样本学习或跨域表征思想。",
    112: "FSD 是这一方向的代表方法，其核心价值在于把 AI 伪造图像识别问题推进到少样本适配与未知生成器场景。",
    113: "LASTED 等工作则从语言引导或对比学习角度增强可迁移表征，为进一步提升泛化能力提供了参考。",
    115: "另一条重要路线关注判别依据的可靠性，即避免模型把“真实图像特征”误当作决定性证据。",
    116: "Stay-Positive 强调通过抑制真实图像特征依赖来提升鲁棒性，其意义不只在于提高指标，更在于纠正错误取证路径。",
    117: "GenDet 等工作也说明，面对新生成器时，输出分布不稳定和伪相关问题往往同时存在。",
    119: "LVLM 路线尝试利用视觉编码、语言理解和跨模态推理能力，对复杂场景中的语义不一致与局部异常进行辅助分析。",
    120: "这类方法更适合处理边界样本、冲突样本和难以用单一分数解释的复杂案例。",
    121: "但 LVLM 的输出受提示词和上下文影响较大，当前更适合作为解释接口，而非直接替代专门训练的检测器。",
    123: "高质量数据集和统一评测基准是比较不同方法的前提。",
    124: "GenImage 是当前极具代表性的公开基准之一，覆盖多类生成器，适合同时观察固定判别能力和跨生成器差异。",
    125: "因此，本文以 GenImage 作为统一实验基础。",
    127: "基于上述研究现状，本文将 FSD、Stay-Positive 与 LVLM 视为解决不同层面问题的三条互补路线，而非简单竞争关系。",
    128: "本文的切入点是基于公开资源建立可信的最小复现基线，并围绕正式结果讨论后续可执行的改进路径。",
    130: "综上，现有研究已经从单纯追求高精度扩展到同时关注泛化能力、判别依据稳定性和复杂场景解释能力，本文的研究设计正建立在这一背景之上。",
    133: "FSD 在本文中承担“泛化检测主干基线”的角色，其价值在于对应未知生成器和少样本适配问题。",
    134: "选择 FSD 作为主线，不是因为其当前指标最高，而是因为其结构更适合支撑后续面向开放场景的研究。",
    136: "Stay-Positive 在本文中被视为“鲁棒判别增强基线”，其重点在于减少模型对真实图像特征的错误依赖。",
    137: "这一思想对本文很重要，因为泛化能力与判别依据稳定性并不是替代关系，而是需要同时考虑的两个维度。",
    139: "LVLM 被定位为“复杂场景辅助分析与解释接口”，用于补充传统判别模型难以说明的边界样本和冲突样本。",
    140: "因此，LVLM 在本文中的主要价值是提供高层语义视角，而不是与 FSD、Stay-Positive 做机械化数值排名。",
    142: "三类方法的关系可以概括为：FSD 提供开放场景主干，Stay-Positive 提供鲁棒判别思路，LVLM 提供复杂案例解释能力。",
    143: "后续若条件允许，可以进一步探索在 FSD 中吸收 Stay-Positive 思想，并用 LVLM 参与难样本分析。",
    145: "本文的研究边界是：当前不宣称完成联合训练，也不宣称已经得到 LVLM 的系统性定量结果，重点仍在可信复现和结果整理。",
    147: "整体方案可概括为三步：完成基线复现；在统一数据上比较方法差异；结合样本级分析提出后续改进方向。",
    149: "本文当前阶段的创新主要体现在研究组织方式和问题聚焦方式上，即把泛化检测、伪相关抑制和复杂场景解释统一纳入同一论文框架。",
    150: "已有实验结果说明，这一方案在资源和时间可控的条件下具有可行性。",
    153: "本文实验主要基于公开数据集 GenImage 展开，并围绕 Midjourney、Stable Diffusion、ADM 与 real 类图像构建最小可运行验证环境。",
    155: "在数据组织上，FSD 直接按目录结构读取数据，Stay-Positive 则通过真假对照 CSV 进行验证，两者因此采用不同输入方式。",
    156: "为便于统一叙述，本文重点围绕 real 与 Midjourney、real 与 Stable Diffusion 两组正式对照任务，以及 ADM 扩展观察展开分析。",
    158: "模型方面，FSD 使用公开实现提供的官方 checkpoint，并完成权重兼容性处理；Stay-Positive 使用公开提供的 Corvi+ 与 Rajan/Ours+ 预训练模型。",
    159: "因此，本文的实现策略是“预训练优先、复现优先”，先确保公开基线在当前环境中稳定运行。",
    161: "FSD 的实验流程包括数据整理、checkpoint 适配、正式评估和 ADM 微调探索。",
    162: "Stay-Positive 的实验流程包括生成 CSV、运行预训练模型、汇总分数并补充 ADM 校准分析。",
    164: "评价指标主要采用 Accuracy 与 AP，前者反映固定阈值下的整体分类正确率，后者反映不同阈值下的综合排序能力。",
    166: "实现过程中主要遇到三类问题：数据组织成本较高、FSD checkpoint 需要兼容性处理，以及部分训练与推理流程对环境配置较敏感。",
    168: "本文的可复现性主要来自公开数据、公开代码和公开权重，同时也明确承认当前实验覆盖范围和测试协议仍存在阶段性边界。",
    169: "因此，本文结果更适合做阶段性比较与研究支撑，而非绝对排名。",
    170: "对 LVLM 而言，当前已明确其接入位置和案例分析入口，但尚未形成统一口径的正式数值评测。",
    176: "FSD 在六类生成器上均具备一定检测能力，其中 GLIDE 表现最好，Stable Diffusion 次之，而 ADM 与 VQDM 相对较弱。这说明 FSD 已具备跨生成器识别能力，但不同生成器之间仍存在明显难度差异。",
    181: "Stay-Positive 在 Midjourney 与 Stable Diffusion 的固定真假对照测试中表现出很强的直接判别能力，两种模型的准确率和 AP 都显著高于 FSD 当前结果。",
    182: "不过，在 ADM 扩展观察中，Rajan/Ours+ 的真实样本准确率为 99.77%，伪造样本准确率仅为 4.73%，整体准确率为 52.25%，而 AP 仍达到 87.58%。这表明其问题更像阈值失配与分数分布漂移，而不是完全失去排序能力。",
    184: "若仅比较当前已获得的正式结果，Stay-Positive 在固定二分类任务中明显优于 FSD；但两类方法的任务定位并不完全一致，因此不宜简单得出“全面更优”的结论。",
    185: "更稳妥的解释是：Stay-Positive 更适合作为高性能判别基线，FSD 更适合作为研究未知生成器泛化与少样本适配的主干方法。",
    186: "ADM 扩展观察进一步提示，高精度闭集表现不能直接外推到所有生成器场景，跨生成器条件下仍需关注阈值校准与输出稳定性。",
    188: "围绕 ADM，本文还完成了从零初始化训练、首轮 checkpoint 微调和第二轮保守微调三组探索。",
    191: "结果显示，从零初始化训练的主要价值在于验证训练链路，而不是带来性能提升；其 ADM 指标明显低于官方基线。",
    192: "基于官方 checkpoint 的首轮微调显著优于随机初始化训练，并把 ADM 结果重新拉回到接近官方基线的水平。",
    193: "第二轮保守微调未进一步改善 ADM，反而略有回落，说明当前任务对微调配置较为敏感。",
    194: "因此，训练探索部分最重要的结论是：强基线初始化是必要的，而继续保守调参并不一定带来收益。",
    196: "从整体结果看，本文得到三点启示：基线复现路径是可行的；固定二分类高分并不能替代泛化能力分析；ADM 是最值得深入研究的方法差异场景。",
    197: "FSD 六类结果的分层现象说明，所谓“跨生成器泛化”并不是均匀成立的，不同生成器之间的难度差异很大。",
    198: "其中，GLIDE 与 Stable Diffusion 的识别相对稳定，而 ADM 与 VQDM 暴露出更明显的挑战性。",
    199: "更重要的是，ADM 的样本级统计把这种差异从“类别差异”推进到了“失败模式差异”：在 3000 个 ADM 假样本中，存在 1241 个校准后 Stay-Positive 与三组 FSD 全冲突的样本，其中主导模式“SP=fake;FSD=real/real/real”达到 1137 个。",
    200: "图 5-1 说明真实样本分数主要集中在 0.36-0.40，伪造样本则更多落在 0.38-0.46，默认阈值 0.5 对 ADM 明显偏高。",
    203: "图 5-2 进一步表明，ADM 场景更值得强调的不是绝对排名，而是不同方法在阈值敏感性和判别依据上的差异。",
    206: "图 5-3 展示的 Top12 代表案例说明，FSD 在 ADM 上的系统性盲区主要集中在伪文本/伪界面、局部结构连接异常、生物体局部真实性不足以及局部修补等细节。",
    210: "尽管本文已经获得较为明确的结果，但仍存在局限：两类方法的覆盖范围与测试协议并不完全一致，ADM 相关证据也主要集中在单一场景。",
    211: "不过，阈值分析已经为 Stay-Positive / ADM 的校准问题提供了较强支持：阈值从 0.5 调整到 0.388818 后，整体准确率可提升到 80.77%。",
    212: "分数分布与冲突模式统计进一步说明，两类方法的差异已不仅是简单阈值问题，而是涉及局部异常证据与整体自然感的不同取证偏好。",
    213: "目前，这种更深层差异主要由案例观察与局部统计支持，仍有待后续更大规模验证。",
    214: "LVLM 当前也尚未形成与前两类模型可直接对比的统一定量结果，因此本文对其角色保持保守表述。",
    215: "总体上，本文在结果解释上坚持“已完成什么就写什么”的原则，即承认 Stay-Positive 的高性能，也保留 FSD 的研究价值，并明确 LVLM 主要承担解释和辅助分析任务。",
    217: "截至目前，本文已经完成多模型框架重构、最小数据子集整理、正式基线复现、ADM 校准分析、样本级冲突统计和训练探索，为后续答辩和定稿提供了稳定基础。",
    220: "在继续推进论文之前，仍需面对几个现实问题：实验覆盖范围有限、不同方法协议不完全一致，以及 LVLM 仍缺少统一定量结果。",
    221: "因此，当前阶段更适合围绕已有结果做深化分析，而不是盲目追加高风险训练。",
    222: "已有三轮训练探索说明，训练链路已经打通，但后续是否继续训练必须以新的结构化证据为前提。",
    224: "基于现有结果，后续总体思路是以 FSD 为主干，吸收 Stay-Positive 的鲁棒判别思想，并把 LVLM 作为复杂样本解释接口纳入整体方案。",
    225: "具体工作包括：继续梳理跨类别差异；围绕误判、分数分布和阈值敏感性做轻量分析；把 ADM 失败模式与案例观察进一步压缩进论文讨论。",
    227: "结合项目进度，下一阶段不再以继续大规模扩类或盲目微调为首要目标，而应优先完成结果收束、问题解释和论文文本固化。",
    229: "从论文写作角度看，较为稳妥的安排是先完成基线复现与结果讨论，再在此基础上补充异常解释和 LVLM 相关分析，必要时最后再评估是否补充极小规模实验。",
    231: "预期成果是形成一份结构完整、结果真实、边界清楚的毕业论文，并为后续围绕 FSD 主干的进一步改进保留明确入口。",
    233: "本文围绕未知生成器泛化、鲁棒判别与复杂样本解释三个方向，对 FSD、Stay-Positive 与 LVLM 开展了基线复现与比较分析。",
    234: "结果表明，FSD 在六类生成器上表现出明显分层，Stay-Positive 在固定真假对照任务中更强，而 ADM 场景则同时暴露出校准问题与稳定的样本级互补性。尤其是 Rajan/Ours+ 在 ADM 上默认阈值下仅得 52.25%，经 0.388818 校准后可提升到 80.77%；同时，3000 个 ADM 假样本中存在 1137 个“Stay-Positive 判 fake 而 FSD 三组均判 real”的主导互补样本。训练探索还说明，基于官方 checkpoint 的微调明显优于随机初始化，但更保守的第二轮微调未继续改善 ADM。",
    235: "综合来看，Stay-Positive 更适合作为高性能判别基线，FSD 更适合作为后续泛化检测研究主干，LVLM 更适合作为复杂样本的语义分析补充。本文已经完成了从基线复现到样本级分析的关键过渡，并为后续最小改动改进实验提供了基础。",
    237: "感谢指导老师在选题、实验推进和论文写作中的持续指导，也感谢公开论文、开源代码、预训练权重与公开数据为本文研究提供基础支持。",
    239: "[1] Wu S, Liu J, Li J, et al. Few-Shot Learner Generalizes Across AI-Generated Image Detection[C]. ICML, 2025.",
    240: "[2] Rajan A S, Lee Y J. Stay-Positive: A Case for Ignoring Real Image Features in Fake Image Detection[C]. ICML, 2025.",
    241: "[3] Zhu M, Chen H, Yan Q, et al. GenImage: A Million-Scale Benchmark for Detecting AI-Generated Image[C]. NeurIPS Datasets and Benchmarks Track, 2023.",
    242: "[4] Wang S Y, Wang O, Zhang R, et al. CNN-Generated Images Are Surprisingly Easy to Spot... for Now[C]. CVPR, 2020.",
    243: "[5] Wu H, Zhou J, Zhang S. Generalizable Synthetic Image Detection via Language-guided Contrastive Learning[EB/OL]. arXiv:2305.13800, 2023.",
    244: "[6] Zhu M, Chen H, Huang M, et al. GenDet: Towards Good Generalizations for AI-Generated Image Detection[EB/OL]. arXiv:2312.08880, 2023.",
}


def char_count(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def set_paragraph_text(paragraph, text: str) -> None:
    """尽量保留段落现有格式，只替换其中的文字内容。"""

    first_run = paragraph.runs[0] if paragraph.runs else None
    run_props = {}
    if first_run is not None:
        run_props = {
            "bold": first_run.bold,
            "italic": first_run.italic,
            "underline": first_run.underline,
            "font_name": first_run.font.name,
            "font_size": first_run.font.size,
        }

    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)

    new_run = paragraph.add_run(text)
    if run_props:
        new_run.bold = run_props["bold"]
        new_run.italic = run_props["italic"]
        new_run.underline = run_props["underline"]
        new_run.font.name = run_props["font_name"]
        new_run.font.size = run_props["font_size"]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="压缩论文 DOCX 内容，尽量不改动既有版式。")
    parser.add_argument("--input", required=True, help="待压缩的 DOCX 路径。")
    parser.add_argument("--backup", help="备份文件路径，默认自动生成。")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = Path(args.input)
    if not input_path.exists():
        raise FileNotFoundError(f"未找到文件：{input_path}")

    backup_path = Path(args.backup) if args.backup else input_path.with_name(f"{input_path.stem}_压缩前备份{input_path.suffix}")
    if not backup_path.exists():
        shutil.copy2(input_path, backup_path)

    doc = Document(str(input_path))
    before = sum(char_count(p.text) for p in doc.paragraphs)

    for idx, text in REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            raise IndexError(f"段落索引超出范围：{idx}")
        set_paragraph_text(doc.paragraphs[idx], text)

    after = sum(char_count(p.text) for p in doc.paragraphs)
    doc.save(str(input_path))

    print(f"backup: {backup_path}")
    print(f"before_chars: {before}")
    print(f"after_chars: {after}")


if __name__ == "__main__":
    main()
