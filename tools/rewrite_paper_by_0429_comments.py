"""根据 4.29 导师修改意见重写论文终稿改写2。

本脚本用于处理根目录下的 ``AI伪造图像识别论文_终稿_改写2.docx``。
它保留原文档的封面和诚信声明，从论文标题页开始重建摘要、目录和正文结构。

修改原则：
1. 按导师意见把第 2 章定位为“已有研究”，第 3 章定位为“本文方案”；
2. 第 5 章压缩为三条主线：多模型基线比较、ADM 困难场景分析、训练探索与联合验证；
3. 第 6 章压缩为“研究总结”和“后续工作展望”两个小节；
4. 保留联合训练结果，但采用谨慎口径，不夸大 LVLM 对主任务性能的提升；
5. 插入必要表格图片和框架图，避免 Word 原生表格在渲染器中跑版。
"""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "AI伪造图像识别论文_终稿_改写2.docx"
FIGURES_DIR = ROOT / "figures"
FRAMEWORK_FIG = FIGURES_DIR / "paper_framework_0429.png"


def load_font(size: int, bold: bool = False):
    """加载中文字体，保证框架图中的中文在 Windows 环境下正常显示。"""

    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


def make_framework_figure() -> None:
    """生成“本文整体研究框架”图片，用于响应导师要求的第 3 章框架图。"""

    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(48, bold=True)
    box_font = load_font(34, bold=True)
    text_font = load_font(27)
    small_font = load_font(24)

    title = "面向泛化检测的多模型协同分析框架"
    draw.text(((width - draw.textlength(title, font=title_font)) / 2, 40), title, fill=(26, 42, 58), font=title_font)

    boxes = [
        (90, 185, 520, 535, "多模型基线比较", ["FSD：未知生成器泛化主干", "Stay-Positive：鲁棒判别基线", "LVLM：语义分析补充"]),
        (685, 185, 1115, 535, "ADM 困难场景分析", ["阈值漂移与分数分布", "样本级冲突统计", "代表案例与局部异常解释"]),
        (1280, 185, 1710, 535, "两阶段联合验证", ["阶段一：FSD + SP 分数", "阶段二：FSD + LVLM 标签", "观察主任务指标与语义 F1"]),
    ]
    colors = [(224, 239, 250), (235, 245, 226), (252, 236, 222)]
    borders = [(57, 100, 143), (86, 130, 70), (179, 105, 54)]

    for idx, (x1, y1, x2, y2, heading, lines) in enumerate(boxes):
        draw.rounded_rectangle([x1, y1, x2, y2], radius=28, fill=colors[idx], outline=borders[idx], width=4)
        draw.text((x1 + 42, y1 + 34), heading, fill=(22, 42, 56), font=box_font)
        y = y1 + 118
        for line in lines:
            draw.ellipse([x1 + 45, y + 12, x1 + 61, y + 28], fill=borders[idx])
            draw.text((x1 + 82, y), line, fill=(35, 35, 35), font=text_font)
            y += 70

    # 画出三条主线之间的流程箭头，突出“比较—分析—验证”的递进关系。
    for start_x, end_x in [(530, 675), (1125, 1270)]:
        y = 355
        draw.line([start_x, y, end_x, y], fill=(70, 82, 94), width=6)
        draw.polygon([(end_x, y), (end_x - 22, y - 14), (end_x - 22, y + 14)], fill=(70, 82, 94))

    bottom = "输出：基线能力边界、ADM 失败模式、联合训练可行性与后续优化方向"
    draw.rounded_rectangle([250, 650, 1550, 790], radius=24, fill=(245, 248, 252), outline=(130, 150, 170), width=3)
    draw.text(((width - draw.textlength(bottom, font=small_font)) / 2, 704), bottom, fill=(36, 50, 65), font=small_font)

    image.save(FRAMEWORK_FIG, quality=95)


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    """统一新增正文的中英文字体，避免新旧段落格式差异过大。"""

    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc: Document, text: str = "", style: str | None = None, align=None):
    """添加段落并设置统一字体。"""

    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_heading(doc: Document, text: str, level: int) -> None:
    """添加章节标题，保持 Word 标题样式，便于后续更新目录。"""

    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, 14 if level == 1 else 12, bold=True)


def add_picture(doc: Document, path: Path, caption: str, width_cm: float = 14.5) -> None:
    """插入图片及图题，所有图片居中显示，避免遮挡正文。"""

    if not path.exists():
        add_paragraph(doc, f"【缺少图片：{path.name}】")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = add_paragraph(doc, caption, style="Caption", align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in cap.runs:
        set_run_font(run, 10)


def remove_after_paragraph(doc: Document, keep_until_index: int) -> None:
    """删除指定段落之后的所有正文块，保留封面和诚信声明。

    这里删除的是 docx 内部 XML 块元素，不涉及文件系统删除，因此不违反文件删除限制。
    """

    body = doc._body._element
    keep_element = doc.paragraphs[keep_until_index]._p
    seen_keep = False
    for child in list(body):
        if child is keep_element:
            seen_keep = True
            continue
        if seen_keep and child.tag.endswith(("}p", "}tbl")):
            body.remove(child)


def count_chinese_chars(doc: Document) -> int:
    """统计文本层中文字符数，用于确认篇幅大致符合终稿要求。"""

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return sum(1 for char in text if "\u4e00" <= char <= "\u9fff")


def add_manual_toc(doc: Document) -> None:
    """按导师修改意见写入手工目录，页码保留为近似值，后续可在 Word 中更新。"""

    add_paragraph(doc, "目  录", align=WD_ALIGN_PARAGRAPH.CENTER)
    toc_lines = [
        "1 绪论\t8",
        "1.1 研究背景\t8",
        "1.2 研究意义\t8",
        "1.3 研究目标\t9",
        "1.4 研究内容\t9",
        "1.5 研究方法\t10",
        "1.6 论文结构安排\t10",
        "2 基础知识与相关工作\t11",
        "2.1 AI伪造图像识别任务与关键问题\t11",
        "2.2 现有方法分类与代表工作\t12",
        "2.3 FSD、Stay-Positive 与 LVLM 方法概述\t14",
        "2.4 本章小结\t15",
        "3 面向泛化检测的多模型协同分析方案\t16",
        "3.1 现有研究的不足与本文切入点\t16",
        "3.2 问题定义与整体思路\t16",
        "3.3 本文整体研究框架\t17",
        "3.4 多模型比较与困难场景分析方案\t18",
        "3.5 两阶段联合验证方案\t19",
        "3.6 本章小结\t20",
        "4 实验设置与设计\t21",
        "4.1 实验设置\t21",
        "4.2 实验研究问题\t22",
        "4.3 实验设计\t22",
        "4.4 本章小结\t24",
        "5 实验结果与分析\t25",
        "5.1 多模型基线结果与对比分析\t25",
        "5.2 ADM困难场景结果分析\t27",
        "5.3 训练探索与联合验证结果分析\t29",
        "5.4 总体讨论与局限性分析\t32",
        "5.5 本章小结\t33",
        "6 总结与展望\t34",
        "6.1 研究总结\t34",
        "6.2 后续工作展望\t35",
        "参考文献\t36",
        "致谢\t37",
    ]
    for line in toc_lines:
        add_paragraph(doc, line, style="toc 1" if not line[0].isdigit() or "." not in line.split()[0] else "toc 2")


def build_document() -> None:
    """执行整篇论文的结构化改写。"""

    make_framework_figure()
    doc = Document(DOCX_PATH)

    # 保留封面和诚信声明，删除旧摘要、旧目录和旧正文。
    remove_after_paragraph(doc, 12)

    # 标题与摘要。
    add_paragraph(doc, "面向泛化检测的AI伪造图像多模型识别研究", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "摘  要", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        "随着扩散模型和多模态生成技术的发展，AI伪造图像在视觉真实性、内容复杂度和生成方式多样性方面不断提升，传统检测方法面临未知生成器下泛化能力不足、判别依据不稳定以及复杂样本难以解释等问题。针对上述问题，本文围绕 AI 伪造图像识别任务，构建了一个面向泛化检测的多模型研究框架，并对 FSD、Stay-Positive 和 LVLM 三类代表性方法进行了实验研究。首先，在公开代码、预训练权重和 GenImage 数据集基础上，完成了 FSD 与 Stay-Positive 的基线复现，并比较了不同方法在固定真假对照任务和跨生成器场景下的性能差异。其次，以 ADM 为重点困难场景，从阈值漂移、分数分布和样本级冲突等方面分析了不同方法的失败模式与互补特征。最后，设计并实施了两阶段联合验证：第一阶段将 Stay-Positive 的离线分数引入 FSD 主干训练流程，第二阶段进一步加入 LVLM 结构化语义标签和轻量辅助头，检验语义监督是否能够进入训练计算图并形成可量化输出。实验结果表明，三类方法在任务定位和适用场景上各具特点：Stay-Positive 更适合固定真假对照下的高性能判别，FSD 更适合作为未知生成器场景下的研究主干，LVLM 能为复杂样本提供语义分析支持，并具备作为轻量辅助监督信号接入训练流程的可行性。本文形成了一个从多模型基线比较、困难场景分析到联合验证的完整研究过程，可为后续 AI 伪造图像识别中的泛化检测、判别校准和视觉语义协同研究提供参考。"
    )
    add_paragraph(doc, "关键词：AI伪造图像识别；泛化检测；多模型分析；鲁棒判别；语义辅助监督")
    add_paragraph(doc, "Research on Multi-Model AI-Generated Image Detection for Generalization-Oriented Analysis", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Abstract", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        "With the rapid development of diffusion models and multimodal generation systems, AI-generated images have become increasingly realistic, diverse, and difficult to explain. Conventional detection methods therefore face challenges in generalization to unseen generators, stability of discriminative evidence, and interpretation of complex samples. To address these issues, this thesis constructs a multi-model research framework for AI-generated image detection and studies three representative methods, namely FSD, Stay-Positive, and LVLM. Based on public code, pretrained checkpoints, and the GenImage dataset, FSD and Stay-Positive are first reproduced and compared under paired binary settings and cross-generator scenarios. Then ADM is selected as a difficult case for deeper analysis, where threshold drift, score distribution, and sample-level conflicts are examined. Finally, a two-stage joint validation is designed. The first stage introduces offline Stay-Positive scores into the FSD training process, while the second stage further incorporates structured LVLM semantic labels and a lightweight auxiliary head. Experimental results show that Stay-Positive is suitable for high-performance discrimination in fixed real-fake settings, FSD is more suitable as a research backbone for unseen-generator detection, and LVLM can provide semantic support for complex samples and can be connected to the training process as a lightweight auxiliary supervision signal. This work forms a complete research process from baseline comparison and difficult-case analysis to joint validation, providing a reference for generalization detection, discriminative calibration, and visual-semantic collaboration in AI-generated image detection."
    )
    add_paragraph(doc, "Keywords: AI-generated image detection; generalization detection; multi-model analysis; robust discrimination; semantic auxiliary supervision")
    add_manual_toc(doc)

    # 第 1 章。
    add_heading(doc, "绪论", 1)
    add_heading(doc, "研究背景", 2)
    add_paragraph(doc, "近年来，GAN、扩散模型和多模态生成系统持续提升伪造图像的真实感，图像检测任务已经从寻找明显伪迹，转向面对高仿真、跨模型和跨分布的开放场景。早期伪造图像常常带有明显纹理异常、边界不自然或全局结构不协调等缺陷，传统检测方法可以依靠低层统计特征取得较好效果。但随着扩散模型和商业生成工具普及，许多伪造图像在光照、构图、纹理和语义层面都更接近真实照片，检测模型需要面对更隐蔽、更复杂的判别条件。")
    add_paragraph(doc, "在实际应用中，伪造图像可能进入社交媒体传播、新闻核验、平台内容审核、商业宣传和数字取证等场景。此时检测系统不仅要在固定测试集上取得较高准确率，还需要在生成器未知、图像经过压缩或二次编辑、内容风格不断变化的情况下保持稳定。换言之，AI 伪造图像识别不再只是一个封闭二分类问题，而是与域泛化、鲁棒判别、样本解释和工程可复现性密切相关的综合研究问题。")
    add_paragraph(doc, "当前研究面临的核心困难主要包括三点：一是未知生成器条件下的泛化能力不足，模型在一个生成器上表现良好，不代表能够迁移到新的生成模型；二是判别依据可能不稳定，模型有时会依赖真实图像特征或数据集偏差，而不是伪造图像本身的稳定证据；三是复杂样本难以解释，特别是在伪文本、局部结构异常、局部修补和生物体细节失真等场景中，单一分数往往难以说明模型为何判断错误。")
    add_heading(doc, "研究意义", 2)
    add_paragraph(doc, "从理论意义看，AI伪造图像识别涉及视觉表示学习、域泛化、伪相关抑制和多模态语义理解等问题。对多种代表性方法进行统一复现和比较，可以帮助判断不同技术路线的能力边界，避免只凭单一指标判断模型优劣。特别是在未知生成器场景中，平均准确率并不能充分反映模型是否真正学到了稳定伪造证据，因此需要结合分数分布、阈值校准和样本级冲突来分析。")
    add_paragraph(doc, "从应用意义看，稳定可靠的伪造图像检测能力可服务于媒体审核、新闻核验、平台治理和数字取证。若检测模型只能在理想化闭集条件下工作，那么其在真实部署中的价值会受到限制。因此，围绕泛化检测、鲁棒判别和复杂样本解释开展系统研究，有助于提升检测方法在真实场景中的可用性。")
    add_heading(doc, "研究目标", 2)
    add_paragraph(doc, "本文目标不是提出一个完全成熟的新检测器，而是在公开资源条件下构建一个面向泛化检测的多模型协同分析框架。具体来说，本文希望通过 FSD、Stay-Positive 和 LVLM 三类代表方法，分别观察未知生成器泛化能力、鲁棒判别能力和复杂样本语义解释能力，并进一步检验这些信息能否在联合验证流程中形成可运行、可记录、可评价的研究闭环。")
    add_heading(doc, "研究内容", 2)
    add_paragraph(doc, "本文研究内容围绕三条主线展开。第一条主线是多模型基线比较，即在 GenImage 数据和公开权重基础上复现 FSD 与 Stay-Positive 的代表性结果，比较其在跨生成器场景和固定真假对照任务中的表现差异。第二条主线是 ADM 困难场景分析，即以 ADM 为重点对象，分析 Stay-Positive 在该场景下的阈值漂移、分数分布和样本级冲突，并结合代表样本讨论不同方法的互补性。第三条主线是两阶段联合验证，即先将 Stay-Positive 离线分数接入 FSD 主干训练，再进一步引入 LVLM 结构化语义标签和轻量辅助头，验证语义监督是否能够进入训练计算图并产生量化输出。")
    add_heading(doc, "研究方法", 2)
    add_paragraph(doc, "本文采用文献分析法、代码复现法、结果对比法、样本级统计法和最小量化训练法。文献分析法用于梳理现有检测方法的发展脉络和不足；代码复现法用于确认公开方法在当前环境中的可运行性；结果对比法用于比较不同模型的指标表现；样本级统计法用于分析 ADM 场景下的阈值漂移和冲突样本；最小量化训练法用于在有限资源下验证两阶段联合方案的可行性。")
    add_heading(doc, "论文结构安排", 2)
    add_paragraph(doc, "全文共分为六章。第1章介绍研究背景、意义、目标、内容和方法；第2章梳理 AI 伪造图像识别相关基础知识与已有研究；第3章提出本文面向泛化检测的多模型协同分析方案；第4章说明实验设置、研究问题和实验设计；第5章围绕多模型基线比较、ADM 困难场景分析以及训练探索与联合验证展开结果分析；第6章总结全文工作并给出后续展望。")

    # 第 2 章。
    add_heading(doc, "基础知识与相关工作", 1)
    add_heading(doc, "AI伪造图像识别任务与关键问题", 2)
    add_paragraph(doc, "AI伪造图像识别任务通常被建模为真实图像与生成图像之间的二分类问题，但随着生成模型快速迭代，该任务已经明显超出普通二分类范畴。检测模型不仅要区分当前数据集中的真假图像，还要在未知生成器、未知后处理方式和不同图像内容分布下保持稳定表现。因此，泛化性、鲁棒性和可解释性成为评价检测方法的重要维度。")
    add_paragraph(doc, "未知生成器泛化是当前任务中的关键难点。许多模型在训练生成器或相近生成器上表现较好，但面对新的扩散模型、不同采样策略或不同提示词生成的图像时，性能可能明显下降。判别依据稳定性也是重要问题，如果模型主要依赖数据集背景、压缩痕迹或真实样本偏差，就可能在跨域测试中失效。复杂样本解释问题则体现在模型输出分数无法说明具体异常来源，导致结果难以被人工审核或数字取证流程使用。")
    add_heading(doc, "现有方法分类与代表工作", 2)
    add_heading(doc, "面向固定生成器的检测方法", 3)
    add_paragraph(doc, "面向固定生成器的检测方法通常假设训练集和测试集中的生成器类型相对接近，因此可以通过卷积神经网络、频域特征或局部纹理统计学习真实图像与生成图像之间的差异。这类方法在闭集条件下往往能够取得较高性能，但当生成器发生变化时，学习到的特征可能不再稳定。")
    add_heading(doc, "面向未知生成器的泛化检测方法", 3)
    add_paragraph(doc, "面向未知生成器的泛化检测方法更强调跨模型迁移能力。相关研究通常尝试学习更抽象、更少依赖具体生成器伪迹的表示，或者通过少样本学习、跨域训练和数据增强等方式提升模型对新生成器的适应能力。FSD 就属于这一方向中的代表方法，其核心价值在于为未知生成器检测提供了较适合继续扩展的研究主干。")
    add_heading(doc, "面向鲁棒判别与伪相关抑制的方法", 3)
    add_paragraph(doc, "部分研究关注模型是否错误利用真实图像特征或数据集偏差进行判断。Stay-Positive 的思想正是从这一角度出发，强调忽略真实图像中的干扰特征，促使模型更多关注伪造样本本身的判别证据。这类方法在固定真假对照任务中通常表现较强，但在跨生成器和阈值迁移场景中仍需要进一步校准。")
    add_heading(doc, "面向视觉语义分析的多模态方法", 3)
    add_paragraph(doc, "随着大型视觉语言模型的发展，LVLM 被逐渐用于复杂图像内容理解和异常描述。与纯分数检测器不同，LVLM 更擅长用自然语言描述局部结构不协调、伪文本不自然、物体关系异常等现象。虽然这类模型当前不适合直接替代专门检测器，但它可以为难样本分析和语义辅助监督提供新的信息来源。")
    add_heading(doc, "FSD、Stay-Positive 与 LVLM 方法概述", 2)
    add_paragraph(doc, "FSD、Stay-Positive 和 LVLM 分别对应本文关注的三个层面。FSD 面向未知生成器泛化，适合作为开放场景研究主干；Stay-Positive 面向鲁棒判别和伪相关抑制，适合作为固定真假对照下的高性能判别基线；LVLM 面向视觉语义理解，适合用于复杂样本解释和轻量语义辅助监督。三者不是简单竞争关系，而是从不同角度补充 AI 伪造图像识别任务。")
    add_heading(doc, "本章小结", 2)
    add_paragraph(doc, "本章从任务特点、关键问题和代表方法三个角度梳理了相关研究。已有研究表明，AI伪造图像识别既需要关注闭集检测精度，也需要关注未知生成器泛化、判别依据稳定性和复杂样本解释能力。基于此，本文在后续章节中将三类代表方法纳入同一研究框架，展开多模型比较、困难场景分析和联合验证。")

    # 第 3 章。
    add_heading(doc, "面向泛化检测的多模型协同分析方案", 1)
    add_heading(doc, "现有研究的不足与本文切入点", 2)
    add_paragraph(doc, "现有研究虽然在不同数据集和任务设定下取得了较好结果，但仍存在三方面不足。第一，许多方法更关注单一测试协议下的性能，缺少对未知生成器泛化能力的持续分析。第二，部分模型在闭集任务中表现较好，但其判别依据是否稳定仍不清楚。第三，复杂样本的错误原因往往难以仅凭分数解释，缺少与语义观察相结合的样本级分析。")
    add_paragraph(doc, "本文的切入点是把多模型比较、困难场景分析和联合验证结合起来。通过比较 FSD 与 Stay-Positive，可以观察泛化检测主干与鲁棒判别基线之间的差异；通过聚焦 ADM，可以分析阈值漂移和冲突样本；通过引入 LVLM 结构化标签，可以进一步检验语义信息是否具备进入训练流程的可行性。")
    add_heading(doc, "问题定义与整体思路", 2)
    add_paragraph(doc, "本文研究的问题可以概括为：在公开数据和有限训练资源条件下，如何组织 FSD、Stay-Positive 和 LVLM 三类方法，使其服务于泛化检测、鲁棒判别和复杂样本解释这三个目标。本文不将三类方法简单视为排名对象，而是强调其任务定位差异和互补关系。")
    add_paragraph(doc, "整体思路是先比较，再分析，最后验证。先通过基线复现获得多模型基本表现；再以 ADM 为困难场景深入分析失败模式；最后通过两阶段联合验证观察不同信号是否能够被接入同一训练流程。该思路既能避免只做结果罗列，也能避免在证据不足时直接宣称提出成熟新模型。")
    add_heading(doc, "本文整体研究框架", 2)
    add_paragraph(doc, "本文整体研究框架如图3-1所示。框架由三部分组成：第一部分为多模型基线比较，用于明确 FSD 与 Stay-Positive 的初始能力边界；第二部分为 ADM 困难场景分析，用于解释不同方法在同一批样本上的分歧来源；第三部分为两阶段联合验证，用于检验 Stay-Positive 分数和 LVLM 语义标签是否能够接入 FSD 主干训练流程。")
    add_picture(doc, FRAMEWORK_FIG, "图3-1 本文整体研究框架", 14.5)
    add_heading(doc, "多模型比较与困难场景分析方案", 2)
    add_heading(doc, "多模型基线比较方案", 3)
    add_paragraph(doc, "多模型基线比较方案主要围绕 FSD 和 Stay-Positive 展开。FSD 采用公开预训练权重，在 Midjourney、Stable Diffusion、ADM、BigGAN、GLIDE 和 VQDM 等生成器上进行评估；Stay-Positive 则在 Midjourney、Stable Diffusion 等固定真假对照任务中进行复现。比较时不直接用单一指标判断谁全面更优，而是结合任务协议说明不同方法的适用范围。")
    add_heading(doc, "ADM困难场景分析方案", 3)
    add_paragraph(doc, "ADM 被选为重点困难场景，是因为它同时暴露了阈值漂移、分数分布重叠和样本级冲突等问题。本文从三个角度展开分析：首先观察默认阈值与校准阈值下的性能变化；其次绘制真实样本与伪造样本的分数分布；最后筛选 Stay-Positive 与 FSD 判断不一致的代表样本，讨论不同方法关注的视觉证据是否存在互补。")
    add_heading(doc, "两阶段联合验证方案", 2)
    add_heading(doc, "第一阶段：FSD 与 Stay-Positive 的联合验证", 3)
    add_paragraph(doc, "第一阶段以 FSD 为主干，引入 Stay-Positive 的离线分数作为外部判别参考。该阶段的目的不是立即追求显著性能提升，而是验证离线分数能否稳定进入训练元数据和闭集训练流程，并观察主任务指标是否保持稳定。训练完成后，重点记录 ADM、SD 和 Midjourney 的 Accuracy 与 AP，并结合日志判断 SP 约束是否形成有效监督。")
    add_heading(doc, "第二阶段：FSD 与 LVLM 的联合验证", 3)
    add_paragraph(doc, "第二阶段进一步引入 LVLM 结构化语义标签和轻量辅助头。该阶段关注的问题是：LVLM 是否能够从后验解释工具转变为可训练的辅助监督信号。训练过程中重点观察 valid_lvlm_samples、lvlm_loss、steps_with_valid_lvlm 等日志信号，并在 ADM 上记录 LVLM 辅助头的 F1 指标。")
    add_heading(doc, "本章小结", 2)
    add_paragraph(doc, "本章提出了面向泛化检测的多模型协同分析方案。该方案不是重复介绍已有方法，而是围绕本文研究目标组织三条主线：多模型基线比较用于确定能力边界，ADM 困难场景分析用于解释失败模式，两阶段联合验证用于检验不同模型信号的可接入性。")

    # 第 4 章。
    add_heading(doc, "实验设置与设计", 1)
    add_heading(doc, "实验设置", 2)
    add_heading(doc, "实验环境", 3)
    add_paragraph(doc, "本文实验主要在本地 Windows 环境和服务器 Slurm 环境中完成。本地环境用于数据整理、结果汇总、文档生成和部分轻量分析；服务器环境用于需要 GPU 的模型评估和训练任务。Python 环境采用项目本地虚拟环境，主要依赖 PyTorch、pandas、Pillow、matplotlib 和 python-docx 等工具。")
    add_heading(doc, "数据集与数据组织方式", 3)
    add_paragraph(doc, "实验数据主要来自 GenImage 数据集，涉及真实图像以及 Midjourney、Stable Diffusion、ADM、BigGAN、GLIDE、VQDM 等生成器类别。不同方法的数据输入形式并不完全一致，因此本文在实验设计中对 FSD、Stay-Positive 和联合训练分别整理了对应的数据列表、样本路径和元数据文件。对于不可读图像，训练与评估脚本记录警告并跳过，以保证流程能够继续执行。")
    add_heading(doc, "评测基准与评价指标", 3)
    add_paragraph(doc, "本文主要采用 Accuracy 和 AP 作为主任务指标。Accuracy 反映固定阈值下的分类正确率，AP 反映不同阈值下的排序能力。在第二阶段联合验证中，本文额外记录 LVLM 辅助头在 ADM 上的多标签 F1，用于判断语义标签是否真实进入训练计算图并形成可量化输出。")
    add_heading(doc, "实验研究问题", 2)
    add_paragraph(doc, "本文实验围绕四个研究问题展开：第一，FSD 与 Stay-Positive 在公开数据和预训练权重条件下能够复现出怎样的基线表现；第二，ADM 场景中 Stay-Positive 的低默认准确率是否主要来自阈值漂移；第三，FSD 与 Stay-Positive 在 ADM 样本级判断上是否存在稳定冲突和互补；第四，Stay-Positive 分数与 LVLM 语义标签能否通过两阶段方式接入 FSD 主干训练流程。")
    add_heading(doc, "实验设计", 2)
    add_heading(doc, "多模型基线复现实验设计", 3)
    add_paragraph(doc, "多模型基线复现实验首先运行 FSD 公开预训练模型，并在六类生成器上记录 Accuracy 和 AP；随后运行 Stay-Positive 预训练模型，在固定真假对照任务中记录对应指标。该实验主要用于获得多模型比较的基础数据，而不是直接宣称不同协议下方法的绝对优劣。")
    add_heading(doc, "ADM困难场景分析实验设计", 3)
    add_paragraph(doc, "ADM 困难场景分析包括阈值校准、分数分布统计和冲突样本筛选三部分。阈值校准用于判断默认阈值是否适合 ADM；分数分布用于观察真实样本和伪造样本的重叠区域；冲突样本筛选用于找到 Stay-Positive 与 FSD 判断差异最明显的代表样本，并结合图像内容分析可能的失败模式。")
    add_heading(doc, "两阶段联合验证实验设计", 3)
    add_paragraph(doc, "两阶段联合验证均以 FSD 为主干。第一阶段在训练元数据中加入 Stay-Positive 离线分数，观察该约束是否形成有效监督以及闭集指标是否稳定。第二阶段在元数据中加入 LVLM 结构化标签，并通过轻量辅助头输出语义 F1。两个阶段均记录训练日志、checkpoint、闭集评估结果和关键调试信号。")
    add_heading(doc, "可复现性与边界说明", 3)
    add_paragraph(doc, "本文实验建立在公开代码、公开权重和公开数据之上，关键结果均通过日志、CSV、Markdown 汇总文件和 Word 图表进行保存，具有较好的可追溯性。同时，本文也明确承认其边界：不同方法的原始协议并不完全一致，联合训练属于阶段性最小量化验证，LVLM 标签规模仍有限，因此结论应聚焦于可运行性、互补性和改进方向，而不是夸大为最终性能上限。")
    add_heading(doc, "本章小结", 2)
    add_paragraph(doc, "本章说明了实验环境、数据组织、评价指标和三类实验设计。通过先明确研究问题，再安排基线复现、ADM 分析和两阶段联合验证，本文第5章的结果分析能够与研究目标形成对应关系。")

    # 第 5 章。
    add_heading(doc, "实验结果与分析", 1)
    add_heading(doc, "多模型基线结果与对比分析", 2)
    add_heading(doc, "FSD 基线结果分析", 3)
    add_paragraph(doc, "本文首先完成 FSD 预训练模型在六类生成器上的基线评估。结果如表5-1所示，FSD 在 GLIDE 和 Stable Diffusion 上表现较好，在 ADM 和 VQDM 上相对较弱。这说明 FSD 已具备一定跨生成器检测能力，但不同生成器之间存在明显难度差异，所谓泛化并不是均匀成立的。")
    add_picture(doc, FIGURES_DIR / "table_5_1_fsd_results_clean.png", "表5-1 FSD 基线实验结果", 14.2)
    add_heading(doc, "Stay-Positive 基线结果分析", 3)
    add_paragraph(doc, "Stay-Positive 在固定真假对照任务中表现较强。表5-2显示，其在 Midjourney 和 Stable Diffusion 上均取得很高 Accuracy 与 AP，说明在训练测试协议相近、真假对照关系清晰时，围绕鲁棒判别边界优化的方法具有明显优势。")
    add_picture(doc, FIGURES_DIR / "table_5_2_stay_positive_results_clean.png", "表5-2 Stay-Positive 基线实验结果", 14.2)
    add_heading(doc, "多模型基线结果对比", 3)
    add_paragraph(doc, "从基线结果看，Stay-Positive 更适合固定真假对照下的高性能判别，FSD 更适合作为未知生成器泛化研究主干。二者任务协议和输入形式不同，因此不能简单用单一指标得出全面优劣结论。更合理的解释是：二者分别揭示了闭集判别能力和开放场景泛化能力两个层面的信息。")
    add_heading(doc, "ADM困难场景结果分析", 2)
    add_heading(doc, "阈值漂移分析", 3)
    add_paragraph(doc, "在 ADM 扩展观察中，Stay-Positive 的 Rajan/Ours+ 默认阈值下真实样本准确率为 99.77%，伪造样本准确率仅为 4.73%，整体准确率为 52.25%，但 AP 仍达到 87.58%。这说明其问题并非完全失去排序能力，而更像是默认阈值在 ADM 上不适用。经过阈值校准后，整体准确率可提升到 80.77%，进一步支持了阈值漂移解释。")
    add_picture(doc, FIGURES_DIR / "stay_positive_adm_score_bins.png", "图5-1 Stay-Positive 在 ADM 上的分数分布", 13.5)
    add_picture(doc, FIGURES_DIR / "adm_method_comparison.png", "图5-2 ADM 方法比较图", 13.5)
    add_heading(doc, "冲突样本与代表案例分析", 3)
    add_paragraph(doc, "ADM 样本级统计显示，在 3000 个 ADM 假样本中，存在大量 Stay-Positive 与 FSD 判断不一致的样本，其中主导模式是 Stay-Positive 判断为 fake，而多组 FSD 判断为 real。这类冲突表明，不同方法可能关注到不同证据：Stay-Positive 对部分局部异常和校准后的伪造线索更敏感，而 FSD 在整体自然感较强的图像上更容易出现漏判。")
    add_picture(doc, FIGURES_DIR / "adm_core_conflict_top12_montage.png", "图5-3 ADM 核心冲突代表案例", 14.5)
    add_heading(doc, "训练探索与联合验证结果分析", 2)
    add_heading(doc, "FSD 训练探索结果分析", 3)
    add_paragraph(doc, "围绕 ADM，本文完成了从零初始化训练、基于官方 checkpoint 的首轮微调以及第二轮保守微调。结果如表5-3所示，从零训练主要用于验证训练链路；首轮基于官方 checkpoint 的微调能将 ADM 结果拉回到接近官方基线的水平；第二轮保守微调没有进一步改善 ADM，说明该任务对学习率、训练步数和样本质量较为敏感。")
    add_picture(doc, FIGURES_DIR / "table_5_3_fsd_training_exploration_clean.png", "表5-3 FSD 训练探索结果", 15.0)
    add_heading(doc, "第一阶段联合验证结果分析", 3)
    add_paragraph(doc, "第一阶段 FSD + Stay-Positive 训练完成 10000 step，并在 ADM、SD、Midjourney 三类闭集评估上分别取得 95.50%/97.50%、95.34%/97.90%、87.00%/90.47% 的 Accuracy/AP。从主任务指标看，该结果与同协议 FSD-only 基线基本一致，说明第一阶段主要验证了训练流程可运行，而不是证明 Stay-Positive 分数已经带来显著增益。结合日志中 SP 有效样本不足的现象，后续仍需要改进分数对齐、阈值筛选和采样策略。")
    add_heading(doc, "第二阶段联合验证结果分析", 3)
    add_paragraph(doc, "第二阶段最小量化版进一步引入 LVLM 结构化语义标签。服务器作业 17862387 完成 10000 step 训练，日志显示训练过程中多次出现非零 valid_lvlm_samples 和 lvlm_loss，最终 steps_with_valid_lvlm 达到 6094，avg_valid_lvlm_samples_per_step 为 1.6343。这说明 LVLM 辅助头并非停留在论文方案或后验解释层面，而是已经参与训练计算。")
    add_paragraph(doc, "最终 10000 step 评估中，ADM、SD、Midjourney 的 Accuracy/AP 分别为 95.22%/97.45%、95.33%/97.72%、86.97%/90.38%，ADM 上额外输出 LVLM F1=0.1778。表5-4和图5-4显示，第二阶段主任务指标总体保持稳定，但没有显著超过第一阶段。因此，本文将第二阶段定位为最小量化验证，即证明语义监督可以进入训练流程，而不是证明其已经显著提升主任务性能。")
    add_picture(doc, FIGURES_DIR / "table_5_4_joint_training_results_clean.png", "表5-4 两阶段联合训练结果对比", 15.0)
    add_picture(doc, FIGURES_DIR / "joint_stage2_min_comparison.png", "图5-4 两阶段联合训练闭集指标对比", 13.5)
    add_heading(doc, "总体讨论与局限性分析", 2)
    add_paragraph(doc, "综合来看，多模型基线比较说明不同方法具有不同能力边界：FSD 更适合支撑未知生成器泛化研究，Stay-Positive 在固定真假对照任务中具有更强判别能力，LVLM 则能为复杂样本提供语义分析和轻量辅助监督。ADM 困难场景分析进一步说明，模型失败并不只来自整体性能不足，还可能来自阈值漂移、分数分布重叠和局部异常证据利用不足。两阶段联合验证则说明，三类方法的信息已经能够在一个训练框架中形成初步闭环。")
    add_paragraph(doc, "本文仍存在局限。首先，不同方法的原始测试协议并不完全一致，因此结果比较主要用于说明方法定位和能力边界，而不是绝对排名。其次，ADM 分析虽然具有代表性，但仍集中在一个生成器场景上，后续需要扩展到更多生成器。最后，第二阶段 LVLM 标签规模有限，LVLM F1=0.1778 说明语义辅助头仍处于初步学习阶段，不能将其解释为成熟的视觉语义联合检测器。")
    add_heading(doc, "本章小结", 2)
    add_paragraph(doc, "本章围绕三条主线总结实验结果：多模型基线比较明确了 FSD 与 Stay-Positive 的任务差异；ADM 困难场景分析揭示了阈值漂移和样本级冲突；训练探索与联合验证说明强基线初始化的重要性，以及 LVLM 语义监督接入训练流程的可行性。")

    # 第 6 章。
    add_heading(doc, "总结与展望", 1)
    add_heading(doc, "研究总结", 2)
    add_paragraph(doc, "本文围绕 AI 伪造图像识别中的泛化检测、鲁棒判别和复杂样本解释问题，构建了一个多模型协同分析框架，并完成了从基线复现、困难场景分析到两阶段联合验证的研究过程。主要工作包括：复现 FSD 与 Stay-Positive 的代表性结果，分析 ADM 场景下的阈值漂移和样本级冲突，构建 LVLM 结构化语义标签，并完成 FSD + Stay-Positive 与 FSD + LVLM 的两阶段联合验证。")
    add_paragraph(doc, "主要结论包括三点。第一，FSD 在跨生成器检测中具有研究主干价值，但不同生成器之间存在明显性能分层。第二，Stay-Positive 在固定真假对照任务中表现较强，但在 ADM 上存在明显阈值漂移，经过校准后性能显著改善。第三，LVLM 结构化标签已经能够作为轻量辅助监督进入训练流程，并输出 LVLM F1，但当前仍不能证明其显著提升主任务性能。")
    add_paragraph(doc, "本文不足也较为明确：不同方法协议尚未完全统一，ADM 分析仍需扩展到更多生成器，LVLM 标签规模和质量仍有限，联合训练损失权重和采样策略也需要进一步优化。因此，本文更适合作为一个真实、可复核、边界清楚的阶段性研究，而不是最终工程化系统。")
    add_heading(doc, "后续工作展望", 2)
    add_paragraph(doc, "后续工作可以从三个方向展开。首先，扩展 LVLM 结构化标签覆盖范围，增加更多生成器和更多异常类型样本，使语义辅助监督不再集中于少量 ADM 案例。其次，改进 Stay-Positive 离线分数对齐方式，结合阈值校准和置信度筛选提升第一阶段有效监督样本数量。最后，补充多损失权重、不同标签规模和更多生成器类别下的消融实验，观察主任务 Accuracy/AP 与 LVLM F1 是否能够同时改善。")
    add_paragraph(doc, "在应用层面，未来还可以将样本级冲突分析与人工审核流程结合，使检测系统不仅输出真假判断，还能给出局部异常、语义不一致和模型分歧来源等解释信息。这对于新闻核验、平台治理和数字取证等场景具有现实意义。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Wu S, Liu J, Li J, et al. Few-Shot Learner Generalizes Across AI-Generated Image Detection[C]. ICML, 2025.",
        "[2] Rajan A S, Lee Y J. Stay-Positive: A Case for Ignoring Real Image Features in Fake Image Detection[C]. ICML, 2025.",
        "[3] Zhu M, Chen H, Yan Q, et al. GenImage: A Million-Scale Benchmark for Detecting AI-Generated Images[J]. arXiv, 2023.",
        "[4] Wang S Y, Wang O, Zhang R, et al. CNN-Generated Images Are Surprisingly Easy to Spot... for Now[C]. CVPR, 2020.",
        "[5] Wu H, Chen J, Zhang S. Generalizable Synthetic Image Detection via Language-guided Contrastive Learning[EB/OL]. arXiv:2401.16992, 2024.",
        "[6] Zhou L, Chen H, Huang M, et al. GenDet: Towards Good Generalizations for AI-Generated Image Detection[EB/OL]. arXiv:2310.08891, 2023.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)
    add_heading(doc, "致谢", 1)
    add_paragraph(doc, "感谢指导老师在论文选题、实验推进和文本修改过程中给予的持续指导，也感谢公开论文、开源代码、预训练权重与公开数据为本文研究提供支持。")

    doc.save(DOCX_PATH)
    print(f"wrote={DOCX_PATH}")
    print(f"cn_chars={count_chinese_chars(Document(DOCX_PATH))}")
    print(f"inline_shapes={len(Document(DOCX_PATH).inline_shapes)}")


if __name__ == "__main__":
    build_document()
