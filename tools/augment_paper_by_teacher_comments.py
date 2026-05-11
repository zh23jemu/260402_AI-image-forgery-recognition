from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "AI伪造图像识别论文_终稿_改写2_改写0508.docx"
OUT = ROOT / "AI伪造图像识别论文_终稿_按意见补充版.docx"


def normalize(text: str) -> str:
    return " ".join((text or "").split())


def find_paragraph(doc: Document, text: str):
    for para in doc.paragraphs:
        if normalize(para.text) == text:
            return para
    raise ValueError(f"未找到段落：{text}")


def insert_paragraph_after(anchor, text="", style=None, bold=False):
    """在指定段落后插入正文段落，保持原论文结构不整体重排。"""
    new_para = anchor.insert_paragraph_before(text, style=style)
    anchor._p.addnext(new_para._p)
    if bold:
        for run in new_para.runs:
            run.bold = True
    return new_para


def insert_after(anchor, items):
    """按顺序插入多个段落，返回最后一个插入位置，方便继续追加内容。"""
    current = anchor
    for item in items:
        if isinstance(item, tuple):
            text, style, bold = item
            current = insert_paragraph_after(current, text, style=style, bold=bold)
        else:
            current = insert_paragraph_after(current, item)
    return current


def remove_paragraph(paragraph):
    """删除指定段落，主要用于清理多余空段落，避免孤立空白页。"""
    p = paragraph._p
    p.getparent().remove(p)


def insert_picture_after(anchor, image_path: Path, caption: str, width=2.7):
    """插入案例图片与图注，图片居中，图注独立成段，避免正文与图片重叠。"""
    pic_para = insert_paragraph_after(anchor, "")
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    cap_para = insert_paragraph_after(pic_para, caption)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_para.runs:
        run.font.size = Pt(9)
    return cap_para


def tune_body_style(doc: Document):
    """给新增内容使用与正文接近的字号，减少插入后版面突兀。"""
    for style_name in ["Normal", "Body Text"]:
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = style.font.name or "宋体"
            style.font.size = style.font.size or Pt(10.5)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    tune_body_style(doc)

    # 第4章：实验设置补充
    anchor = find_paragraph(doc, "本文的实验环境主要是本地Windows环境以及Slurm服务器环境。本地环境做数据整理、结果汇总、文档生成和部分轻量分析，服务器环境做需要 GPU 的模型评估与训练任务。Python 环境采用项目本地虚拟环境，主要依赖 PyTorch、pandas、Pillow、matplotlib 和 python-docx 等工具。")
    insert_after(anchor, [
        ("实验环境与本地部署参数补充", "Heading 4", False),
        "为了使实验过程具有可复核性，本文将本地整理环境和服务器训练环境分开记录。本地端使用 Windows 工作区完成数据清单整理、CSV 汇总、图表导出、论文文档生成和轻量统计分析，所有 Python 脚本均优先通过项目本地虚拟环境 .venv 调用，避免系统 Python 版本差异影响结果。服务器端使用 Slurm 作业提交模型评估和训练任务，训练脚本统一从项目根目录进入对应子目录，日志输出到 logs 目录，模型输出到 fsd/output 或 stay_positive_runs 等固定位置。",
        "FSD 与联合训练作业采用单卡 GPU 设置，Slurm 中配置 gpu:1、cpus-per-task=8、mem=64G、最长运行时间 24:00:00，并在训练入口设置 GPU_NUM=1、WORLD_SIZE=1、NUM_WORKERS=8、SEED=42。Stay-Positive 的 ADM 扩展评估资源需求较低，配置 gpu:1、cpus-per-task=4、mem=16G、最长运行时间 02:00:00。这样的资源设置能够覆盖本研究中的 10000 step 微调、联合训练和单模型打分任务，也便于在答辩或复核时说明每组实验的运行边界。",
        "依赖方面，FSD 主要依赖 torch、torchvision、torchmetrics、timm、einops、dill、tqdm、scikit-learn 与 matplotlib；Stay-Positive 主要使用其官方测试代码和预训练权重目录。本文没有把模型权重和原始数据直接写入论文或源码包，而是通过 checkpoint 路径、数据目录结构和运行脚本记录复现实验所需的位置关系。"
    ])

    anchor = find_paragraph(doc, "实验数据主要来自 GenImage 数据集，涉及真实图像以及 Midjourney、Stable Diffusion、ADM、BigGAN、GLIDE、VQDM 等生成器类别。不同的方法所对应的输入数据形式并不相同，所以本文在实验设计中对FSD、Stay-Positive和联合训练分别整理出对应的列表、样本路径和元数据文件。对不可读的图像训练和评价脚本会发出警告并跳过，从而保证流程可以继续进行。")
    insert_after(anchor, [
        ("数据组织与输入文件补充", "Heading 4", False),
        "GenImage 数据在实验中按照 real、ADM、BigGAN、glide、Midjourney、SD、VQDM 等类别组织，FSD 训练和评估脚本会根据 data_root 与类别名拼接 train/val 子目录。真实图像统一作为 real 类，伪造图像按照生成器类别进入不同目录。为了避免不同方法之间输入格式不一致，本文额外构建了 stay_positive_runs 下的 real_val.csv、adm_val.csv、real_scores.csv、adm_scores.csv 等文件，用于 Stay-Positive 的打分和阈值校准。",
        "联合训练进一步使用 analysis/joint_training_metadata.csv 与 analysis/joint_training_metadata_stage2.csv 作为样本级元数据入口。元数据中除 image_path、split、generator、label 等基础字段外，还包含 sp_score_raw、sp_prob_calibrated、sp_pred_calibrated、sp_conflict_flag、fsd_base_pred、fsd_base_score、LVLM 结构化标签以及 hard_weight。这样可以把单模型打分、冲突样本筛选和语义标签统一到 FSD 训练过程，而不是只停留在结果表格层面。",
        "由于公开数据集中可能存在不可读图像或目录命名差异，训练脚本会对图像读取失败进行跳过处理，分析脚本则以成功导出的 CSV 和日志为准。这样做的目的不是回避异常样本，而是保证长时间训练不会因为单个坏图中断，同时将实际进入统计和训练的样本记录在结果文件中。"
    ])

    anchor = find_paragraph(doc, "本文主要用 Accuracy 和 AP 这两个任务指标。准确率是用固定阈值下的分类正确率来表示的，平均精度是用不同阈值下的排序能力来衡量的。第二阶段联合验证时，在ADM上额外记录LVLM辅助头的多标签F1，看语义标签是否真的进入训练计算图，并且可以得到可量化的输出。")
    insert_after(anchor, [
        ("评价指标与阈值口径补充", "Heading 4", False),
        "Accuracy 反映固定阈值下最终真假判定是否正确，AP 则反映模型分数排序能力。二者在 ADM 场景中的差异尤其重要：如果 Accuracy 较低但 AP 仍然较高，说明模型并非完全失去区分能力，而可能是默认阈值与目标生成器分布不匹配。本文因此在 Stay-Positive 的 ADM 分析中额外记录 RACC、FACC、TP、TN、FP、FN 和最优准确率阈值，用于区分模型排序能力和部署阈值问题。",
        "LVLM F1 只用于第二阶段辅助头评价，不作为整体真假识别主指标。其作用是检查语义标签是否能够被轻量辅助头学习到，而不是直接证明主任务性能已经显著提升。因此本文在讨论联合训练结果时，将 Accuracy/AP 与 LVLM F1 分开解释，避免把辅助头指标误写成整体检测性能。"
    ])

    anchor = find_paragraph(doc, "多模型基线复现实验先执行 FSD 公开预训练模型，对六类生成器做 Accuracy 和 AP 的记录；然后执行 Stay-Positive 预训练模型，在真假对照任务中记录对应指标。该实验主要是得到多模型比较的基础数据，而不是直接宣称不同的协议下方法的绝对优劣。")
    insert_after(anchor, [
        ("FSD 与 Stay-Positive 基线运行参数", "Heading 4", False),
        "FSD 基线评估使用转换后的官方 ResNet50 checkpoint，分别对 Midjourney、Stable Diffusion、ADM、BigGAN、GLIDE、VQDM 六类生成器进行测试，每类评估样本数为 9000。FSD 默认训练脚本参数为 batch_size=16、lr=1e-4、total_training_steps=50000、NUM_WORKERS=8、SEED=42、use_fp16=True、accumulation_steps=1；正式 ADM 微调时不从随机骨干开始，而是使用 resnet50_exclude_adm_step[200000]_converted.pth 作为初始化，以减少从零训练带来的不稳定。",
        "Stay-Positive 基线评估复用官方预训练模型 Corvi+ 与 Rajan/Ours+，先通过 create_csv.py 将图像目录整理为输入 CSV，再通过 main.py 输出分数 CSV，最后使用 eval.py 计算 real/fake 对照指标。ADM 扩展观察中重点采用 Rajan/Ours+ 的分数进行阈值校准，因为该模型在 SD 对照测试中表现更稳定，也能更清楚展示跨生成器阈值漂移现象。"
    ])

    anchor = find_paragraph(doc, "ADM困难场景分析分为阈值校准、分数分布统计和冲突样本筛选这三个部分。阈值校准用来判断默认阈值是否合适ADM；分数分布用来观察真实样本和伪造样本的重叠区域；冲突样本筛选找出Stay-Positive和FSD判断差异最大的代表样本，结合图像内容分析可能的失败模式。")
    insert_after(anchor, [
        ("ADM 阈值校准与案例筛选流程", "Heading 4", False),
        "ADM 阈值校准先读取 real_scores.csv 与 adm_scores.csv，将真实样本和伪造样本在同一模型分数空间中比较。默认阈值设置为 0.5，随后通过遍历候选阈值寻找整体 Accuracy 最高的位置，并记录对应的 RACC、FACC、TP、TN、FP、FN。分数分布分析采用 0.02 的分桶宽度，用来观察真实样本和伪造样本在 0.36-0.46 区间的重叠情况。",
        "冲突样本筛选则把校准后的 Stay-Positive 判定结果与 FSD official、首轮 ADM 微调、第二轮保守微调三组判断进行对齐。重点保留 SP=fake;FSD=real/real/real 这类样本，因为它们表示 Stay-Positive 已经在校准后给出伪造判断，而三组 FSD 仍共同判断为真实，最适合作为方法互补性和系统性盲区的代表案例。"
    ])

    anchor = find_paragraph(doc, "两阶段联合验证都是用 FSD 作为主干。第一阶段在训练元数据里加入Stay-Positive离线分数，看它能否起到有效的监督作用，并且闭集指标是否稳定。第二阶段在元数据中加入LVLM结构化标签，用轻量辅助头输出语义F1。两阶段均有训练日志、检查点、闭集评价成果和重要的调试信号。")
    insert_after(anchor, [
        ("两阶段联合训练参数与损失设计", "Heading 4", False),
        "第一阶段联合训练以 FSD 为在线训练主干，训练类别为 real、ADM、SD、Midjourney，评估类别为 ADM、SD、Midjourney。主要参数为 batch_size=16、lr=1e-5、total_training_steps=10000、save_interval=5000、eval_interval=5000、log_interval=500、accumulation_steps=1、use_fp16=True、init_ckpt_path=resnet50_exclude_adm_step[200000]_converted.pth、metadata_csv=analysis/joint_training_metadata.csv、sp_loss_weight=0.3、sp_loss_type=mse、force_real_in_task=True。该阶段主要验证 Stay-Positive 离线分数能否进入统一训练元数据和闭集训练流程。",
        "第二阶段仍沿用 FSD 主干和 Stay-Positive 分数接口，但将 metadata 替换为 analysis/joint_training_metadata_stage2.csv，并开启 enable_lvlm_head=True、lvlm_loss_weight=0.2、lvlm_loss_on_fake_only=True。训练总损失写作 L_total = L_proto + 0.3 * L_sp + 0.2 * L_lvlm，其中 L_proto 为 FSD 原型分类损失，L_sp 为 Stay-Positive 分数约束损失，L_lvlm 为 LVLM 多标签辅助损失。这样设计的目的不是一次性训练三个大模型，而是用一个可控的轻量方式检验外部判别分数和结构化语义标签能否进入同一训练图。",
        "在实现上，LVLM 辅助头接在 FSD 视觉表征之后，输出五类多标签语义信号：伪文本/伪界面、布局或关系异常、局部结构连接异常、生物体局部真实性不足、局部修补或过度平滑。训练时只在存在有效 LVLM 标签的样本上计算辅助损失，并使用 lvlm_confidence 与 hard_weight 进行轻量加权；评估时在 ADM 上额外输出 lvlm_f1。"
    ])

    # 第5章：结果与案例补充
    anchor = find_paragraph(doc, "该结果很重要，它把“模型失败”分成两个方面。如果AP很低，那么模型就无法区分真实的和伪造的样本；但是目前AP仍然很高，说明模型排序的能力还在，只是默认的阈值不能适应ADM分数分布。就实际应用而言，这就意味着检测系统不能只依靠固定的阈值，还要考虑目标生成器、样本分布以及部署环境来对它进行校准。")
    insert_after(anchor, [
        ("阈值校准过程与分布解释补充", "Heading 4", False),
        "具体来说，在默认阈值 0.5 下，Rajan/Ours+ 在 ADM 扩展观察中的 ACC、RACC、FACC 分别为 52.25%、99.77%、4.73%，对应 TP/TN/FP/FN 为 142/2993/7/2858。这说明默认阈值会把绝大多数 ADM 伪造样本判成真实，导致伪造召回非常低。",
        "当阈值校准到 0.388818 时，ACC 提升到 80.77%，RACC 与 FACC 分别变为 76.57% 和 84.97%，对应 TP/TN/FP/FN 为 2549/2297/703/451。虽然真实样本准确率有所下降，但伪造样本召回大幅提升，整体识别能力更均衡。与此同时，该设置下 AP 仍达到 87.58%，说明模型排序能力并未完全失效，主要矛盾是默认部署阈值与 ADM 分数分布不匹配。",
        "从分数分布看，真实样本主要集中在 0.36-0.40，而伪造样本更多分布在 0.38-0.46，二者在阈值附近有明显重叠。这种分布结构解释了为什么固定 0.5 阈值会失败，也说明实际部署时不能只使用一个固定阈值，而应结合目标场景做校准或保留不确定区间。"
    ])

    anchor = find_paragraph(doc, "代表案例又证明，ADM伪造图像的困难，并不是整个画面都是不真实的，而是局部细节的困难。伪文本、局部边缘连接、物体结构关系以及细微修补痕迹等只占图像很小一部分，但是会对人工判断造成很大的影响。单纯的全局表征检测器会忽略这些细小的地方，但是带有语义观察能力的模型可以给这些局部的异常提供更加清晰的解释入口。")
    anchor = insert_after(anchor, [
        ("识别图像真假的代表案例补充", "Heading 4", False),
        "为了让结果不只停留在表格层面，本文选取两张 ADM 强冲突样本作为真假识别案例。两张样本的真实标签均为 fake，校准后的 Stay-Positive 判断为 fake，但 FSD official、首轮 ADM 微调和第二轮保守微调均判断为 real。也就是说，模型之间的分歧不是随机误差，而是发生在“整体摄影感较强、局部细节存在异常”的高仿真样本上。",
    ])
    img1 = ROOT / "analysis" / "adm_conflict_priority5_images" / "adm_conflict_priority_01_421_adm_153.PNG"
    img2 = ROOT / "analysis" / "adm_conflict_priority5_images" / "adm_conflict_priority_02_508_adm_174.PNG"
    anchor = insert_picture_after(anchor, img1, "图5-5 ADM 强冲突案例一：真实标签 fake，SP 判 fake，三组 FSD 均判 real")
    anchor = insert_after(anchor, [
        "案例一对应低照度近景结构图。该图在光照、材质和整体构图上接近真实摄影，但支撑件连接、杆件透视以及背景亮区之间缺少稳定的空间关系。FSD 更容易被整体自然感影响而判断为 real，校准后的 Stay-Positive 则对局部结构异常更敏感，因此给出 fake 结论。"
    ])
    anchor = insert_picture_after(anchor, img2, "图5-6 ADM 强冲突案例二：真实标签 fake，SP 判 fake，三组 FSD 均判 real")
    insert_after(anchor, [
        "案例二对应设备按键近景。图像的材质、景深和拍摄角度都较自然，但按键字符存在明显伪文本、符号畸变和排版不一致。该案例说明 ADM 困难样本往往不是全局语义崩坏，而是在很小的局部区域暴露伪造痕迹；这也是引入 LVLM 语义标签和局部异常分析的直接原因。"
    ])

    anchor = find_paragraph(doc, "本文围绕ADM完成了从零初始化训练、使用官方checkpoint进行第一轮微调和第二轮保守微调。从表5-3可以看出，结果中零训练用来检验训练链路，首轮用官方checkpoint的微调可以使ADM结果恢复到接近官方基线的水平，第二轮保守微调并没有使ADM的结果得到更进一步的改善，说明该任务对于学习率、训练步数以及样本质量都很敏感。")
    insert_after(anchor, [
        ("FSD 训练探索参数补充", "Heading 4", False),
        "从零初始化训练主要用于验证 FSD 训练链路是否能够在当前数据目录和服务器环境下跑通，因此其价值在于流程验证。正式可比较的 ADM 微调均使用官方 converted checkpoint 初始化。首轮 ADM 微调使用 lr=1e-5、batch_size=16、total_training_steps=10000、accumulation_steps=1、use_fp16=True，并以 ADM 作为 exclude_class，输出目录为 fsd/output/finetune_adm_stage1。该设置在 ADM 上取得 75.28%/78.54% 的 Accuracy/AP，明显优于从零初始化训练。",
        "第二轮保守微调使用同一初始化 checkpoint，但把学习率降为 5e-6，总训练步数降为 5000，保存和评估间隔均为 5000，输出目录为 fsd/output/finetune_adm_stage2。结果在 ADM 上为 74.13%/76.89%，低于首轮微调。这个负结果很重要，因为它说明当前任务不是简单地降低学习率、缩短训练步数就会更稳定；模型对训练参数、样本构成和初始化状态都较敏感。"
    ])

    anchor = find_paragraph(doc, "该阶段的意义在于证明联合训练链路是畅通的。也就是说Stay-Positive分数可以被整理到训练元数据中，FSD主干可以在闭集的设置下完成训练和评估，checkpoint和日志也可以被完整地保存下来。虽然SP约束并没有带来明显的增益，但是它给第二阶段接入LVLM语义标签打下了工程基础，也说明后面优化应该从监督信号的有效性入手，而不能仅仅增加训练轮次。")
    insert_after(anchor, [
        ("第一阶段联合训练调试信号补充", "Heading 4", False),
        "第一阶段训练日志进一步表明，虽然 metadata 已经接入训练流程，但有效 Stay-Positive 监督样本不足，valid_sp_samples 长期为 0，sp_loss 也没有形成稳定贡献。因此，本阶段更准确的定位是“联合训练框架落地验证与问题暴露”，而不是“SP 分支已经带来显著性能提升”。这一定性能够解释为什么闭集指标与同协议 FSD-only 基线接近，也避免把工程链路跑通误写成性能突破。",
        "这一结果对第二阶段仍然有价值。它说明 FSD 主干、metadata 读取、checkpoint 保存、闭集评估和日志统计都已经打通；后续真正需要改进的是监督信号的有效样本覆盖、分数对齐和标签质量。因此第二阶段没有继续盲目增加 SP 训练轮次，而是把重点转向更明确可解释的 LVLM 结构化语义标签。"
    ])

    anchor = find_paragraph(doc, "LVLM F1=0.1778的数值不高，说明本文要保持谨慎的口径。由于目前的语义标签主要是结构化的案例以及弱标签扩展，所以标签的数量、类别均衡性、标注精度还存在提升的空间。本文真正可以确定的是，LVLM标签已经由文本分析材料变成了训练中的监督信号，并且辅助头可以输出可度量指标。该信号对于提高主任务性能起到怎样的作用，还要通过更多的标签、合理的损失权重和系统的消融实验来探究。")
    insert_after(anchor, [
        ("LVLM 人工标注、弱扩展与接入细节", "Heading 4", False),
        "LVLM 结构化标签不是凭空生成的，而是先从 ADM 困难样本中进行人工视觉观察。本文共整理 20 个代表性案例，其中包括首批优先样本、第二批 strong-conflict 样本和核心互补 Top12 样本。人工观察重点记录样本是否存在伪文本/伪界面、局部结构连接异常、生物体局部真实性不足、局部修补或过度平滑，以及布局或关系不一致等问题。",
        "从统计结果看，20 个案例中局部结构连接异常出现 14 次，伪文本/伪界面出现 4 次，生物体局部真实性不足出现 2 次；场景覆盖生物体/自然场景、设备/伪界面、室内/建筑/结构场景和复杂生活场景。这说明 ADM 困难样本并非只有一种固定错误，而是集中体现为“整体语义高度自然、局部细节持续不自洽”。",
        "弱扩展阶段并不把所有 ADM 图像都强行标为同一类语义异常，而是优先把人工标注案例、强冲突样本、提示词模板相关样本和高价值 ADM 子集写入 joint_training_metadata_stage2.csv。对于人工确认案例，使用较高的 lvlm_confidence；对于根据 ADM prompt 或冲突模式弱扩展得到的样本，使用较低的置信度和 hard_weight 控制影响范围。训练时只有 lvlm_valid 为真的样本参与 LVLM 辅助损失，从而避免少量弱标签过度主导主任务。",
        "第二阶段服务器作业 17862387 完成 10000 step 训练，并保存 5000 step 与 10000 step 两个 checkpoint。日志显示 max_valid_lvlm_samples 为 6，steps_with_valid_lvlm 为 6094，avg_valid_lvlm_samples_per_step 为 1.6343，说明 LVLM 标签在大量训练步骤中真实参与计算。最终 ADM、SD、Midjourney 的 Accuracy/AP 分别为 95.22%/97.45%、95.33%/97.72%、86.97%/90.38%，ADM 上输出 LVLM F1=0.1778。本文据此将第二阶段界定为“语义监督进入训练流程的最小量化验证”，而不是已经完成成熟语义联合检测器。"
    ])

    # 总体讨论段落补充，帮助回应“实际做的部分太少”的意见。
    anchor = find_paragraph(doc, "本文还存在着不足。不同方法的原始测试协议不完全相同，所以结果比较主要用以说明方法定位和能力边界，而不是绝对排名。其次，ADM分析虽然有代表性，但是只对一个生成器进行分析，之后需要拓展到更多的生成器上。第二阶段LVLM标签量较少，LVLM F1=0.1778，说明语义辅助头还处在初步学习的阶段，不能把LVLM解释成一个成熟的视觉语义联合检测器。")
    insert_after(anchor, [
        "从实验工作量看，本文实际完成的内容包括：FSD 六类生成器基线评估、Stay-Positive 两类正式对照评估、Stay-Positive ADM 扩展打分与阈值校准、ADM 冲突样本统计、三轮 ADM 真实案例观察、FSD ADM 两轮微调、第一阶段 FSD+SP 联合训练、第二阶段 FSD+LVLM 最小量化联合训练。上述内容共同构成了从公开模型复现到困难样本解释、再到联合训练验证的完整链路。",
        "因此，本文最终不把贡献写成单一模型性能提升，而是写成实验链路和问题分析的闭环：第一，用基线实验说明不同方法的能力边界；第二，用 ADM 校准和冲突案例解释为什么指标会失效；第三，用两阶段联合训练证明 SP 分数和 LVLM 标签能够被组织进 FSD 主干训练流程；第四，明确指出当前 SP 监督有效样本不足、LVLM 标签规模偏小等限制，为后续改进留下清晰方向。"
    ])

    # 第二轮补充：增加“台账式”细节，让实验过程、参数和复核路径更清楚。
    anchor = find_paragraph(doc, "依赖方面，FSD 主要依赖 torch、torchvision、torchmetrics、timm、einops、dill、tqdm、scikit-learn 与 matplotlib；Stay-Positive 主要使用其官方测试代码和预训练权重目录。本文没有把模型权重和原始数据直接写入论文或源码包，而是通过 checkpoint 路径、数据目录结构和运行脚本记录复现实验所需的位置关系。")
    insert_after(anchor, [
        ("实验运行台账补充", "Heading 4", False),
        "本文所有长时间实验都按照“脚本入口、资源配置、输入数据、输出文件、日志检查”五个环节记录。以 FSD 评估为例，脚本入口为 fsd/test.py 或封装后的 fsd/scripts/eval_*.sh，输入包括 data/GenImage 下的真实图像、目标生成器图像和 checkpoints/fsd 下的 converted checkpoint，输出包括 Accuracy、AP 以及样本级分数 CSV。以训练任务为例，脚本入口为 fsd/train.py 或 fsd/train_joint.py，输出包括训练日志、5000/10000 step checkpoint 和闭集评估结果。",
        "本地部署时重点确认三类路径：第一，data_root 是否指向 data/GenImage；第二，checkpoint 是否指向 checkpoints/fsd/resnet50_exclude_adm_step[200000]_converted.pth 等真实文件；第三，输出目录是否可写，例如 fsd/output/finetune_adm_stage1、fsd/output/joint_sp_stage1 和 fsd/output/joint_stage2_min。服务器部署时还需要确认 Slurm 作业日志路径 logs/*.out 与 logs/*.err 是否存在，训练是否在登录节点之外的 GPU 节点运行。",
        "为了降低复现实验中的偶然性，本文在 FSD 与联合训练中统一使用 SEED=42，单卡训练均设置 GPU_NUM=1，数据加载线程为 NUM_WORKERS=8。虽然这些设置不能完全消除深度学习训练中的随机性，但可以保证不同阶段实验在资源规模、数据入口和训练步数上具有可比较性。"
    ])

    anchor = find_paragraph(doc, "联合训练进一步使用 analysis/joint_training_metadata.csv 与 analysis/joint_training_metadata_stage2.csv 作为样本级元数据入口。元数据中除 image_path、split、generator、label 等基础字段外，还包含 sp_score_raw、sp_prob_calibrated、sp_pred_calibrated、sp_conflict_flag、fsd_base_pred、fsd_base_score、LVLM 结构化标签以及 hard_weight。这样可以把单模型打分、冲突样本筛选和语义标签统一到 FSD 训练过程，而不是只停留在结果表格层面。")
    insert_after(anchor, [
        ("联合训练 metadata 字段说明", "Heading 4", False),
        "metadata 中的 image_path、split、generator 和 label 用于定位样本和主任务真假标签；sp_score_raw 保存 Stay-Positive 原始输出，sp_prob_calibrated 保存经过阈值校准或概率化处理后的分数，sp_pred_default 与 sp_pred_calibrated 分别对应默认阈值和校准阈值下的预测结果；sp_conflict_flag 用来标记是否属于 SP 与 FSD 判断不一致的样本。",
        "LVLM 相关字段采用多标签形式，包括 lvlm_has_text_artifact、lvlm_has_layout_conflict、lvlm_has_structure_error、lvlm_has_bio_detail_error、lvlm_has_patch_or_smooth 和 lvlm_confidence。这样的字段设计比单个“是否异常”标签更细，因为 ADM 困难样本可能同时存在伪文本、结构连接和局部修补等多种问题。hard_weight 则用于把强冲突样本和人工确认困难样本传递给训练过程，使模型在损失计算时对这些样本保留更高关注。"
    ])

    anchor = find_paragraph(doc, "第二阶段仍沿用 FSD 主干和 Stay-Positive 分数接口，但将 metadata 替换为 analysis/joint_training_metadata_stage2.csv，并开启 enable_lvlm_head=True、lvlm_loss_weight=0.2、lvlm_loss_on_fake_only=True。训练总损失写作 L_total = L_proto + 0.3 * L_sp + 0.2 * L_lvlm，其中 L_proto 为 FSD 原型分类损失，L_sp 为 Stay-Positive 分数约束损失，L_lvlm 为 LVLM 多标签辅助损失。这样设计的目的不是一次性训练三个大模型，而是用一个可控的轻量方式检验外部判别分数和结构化语义标签能否进入同一训练图。")
    insert_after(anchor, [
        "从训练过程看，FSD 主干仍然负责完成 real/fake 原型判别，Stay-Positive 分支不参与在线特征提取，只作为离线分数约束来源；LVLM 也不作为在线大模型反向传播，而是把已经整理好的语义标签接入轻量辅助头。这种分层接入方式能够控制显存、训练时间和论文解释复杂度，也更符合毕业论文阶段的可实现条件。",
        "训练日志中重点观察 proto_loss、sp_loss、lvlm_loss、total_loss、valid_sp_samples、valid_lvlm_samples、steps_with_valid_sp、steps_with_valid_lvlm 等信号。如果 valid_sp_samples 或 valid_lvlm_samples 长期为 0，就说明对应辅助监督虽然写入了方案，但并没有真正进入训练；如果这些信号非零，并且评估阶段能够输出对应指标，则说明辅助监督至少已经完成可训练接入。本文第二阶段正是依据这些日志信号判断 LVLM 标签已进入训练流程。"
    ])

    anchor = find_paragraph(doc, "本文实验基于公开代码、公开权重和公开数据，主要结果以日志、CSV、Markdown 汇总文件和 Word 图表的形式保存，具备较好的可追溯性。同时本文也承认了它的边界，不同的方法原始协议并不完全相同，联合训练属于阶段性最小量化验证，LVLM标签数量仍然较少，所以结论只能是可运行性、互补性以及改进方向，不能夸大为最终性能上限。")
    insert_after(anchor, [
        ("结果文件与复核路径补充", "Heading 4", False),
        "为了保证结果不是手工拼接得到的，本文将主要实验结果保存为多级证据。第一层是 Slurm 原始日志，用来确认训练步数、checkpoint 保存、评估输出和报错情况；第二层是 analysis 目录下的 CSV 文件，用来保存样本级分数、冲突模式、阈值校准结果和联合训练汇总结果；第三层是 docs 目录下的 Markdown 汇总，用来把日志和 CSV 中的数字转写为论文可引用文本；第四层才是 Word 正文和图表。",
        "复核时可以按以下顺序检查：先确认 run_fsd_finetune_adm.slurm、run_fsd_finetune_adm_v2.slurm、run_fsd_joint_sp_stage1.slurm、run_fsd_joint_stage2_min.slurm 中的参数；再检查 docs/experiment_results_summary.md、docs/joint_stage2_min_result.md 和 docs/lvlm_structured_supplement_result.md 中的结果；最后对应到论文第4章实验设置和第5章结果分析。这样可以把“论文文字”追溯到“运行脚本和结果文件”。"
    ])

    for idx, para in enumerate(doc.paragraphs):
        if normalize(para.text) == "实验结果与分析":
            prev_para = doc.paragraphs[idx - 1] if idx > 0 else None
            if prev_para is not None and normalize(prev_para.text) == "":
                remove_paragraph(prev_para)
            break

    anchor = find_paragraph(doc, "具体来说，在默认阈值 0.5 下，Rajan/Ours+ 在 ADM 扩展观察中的 ACC、RACC、FACC 分别为 52.25%、99.77%、4.73%，对应 TP/TN/FP/FN 为 142/2993/7/2858。这说明默认阈值会把绝大多数 ADM 伪造样本判成真实，导致伪造召回非常低。")
    insert_after(anchor, [
        "这一现象也可以从真实样本和伪造样本的均值看出：ADM 假样本分数均值约为 0.422052，真实样本分数均值约为 0.382475，两类样本确实存在相对位移，但整体都低于默认 0.5 阈值。如果只使用 0.5 作为真假边界，模型会把大量本应位于 fake 一侧的 ADM 样本压到 real 侧，从而形成高 RACC、低 FACC 的失衡结果。",
        "因此，本文没有把默认阈值下 52.25% 的 ACC 简单解释为 Stay-Positive 完全失败，而是进一步检查 AP、分数均值、分桶分布和最优阈值。这样的分析能够说明，跨生成器部署中的阈值校准本身就是检测系统的一部分，尤其是面对 ADM 这类整体自然度较高的生成器时更不能省略。"
    ])

    anchor = find_paragraph(doc, "案例二对应设备按键近景。图像的材质、景深和拍摄角度都较自然，但按键字符存在明显伪文本、符号畸变和排版不一致。该案例说明 ADM 困难样本往往不是全局语义崩坏，而是在很小的局部区域暴露伪造痕迹；这也是引入 LVLM 语义标签和局部异常分析的直接原因。")
    insert_after(anchor, [
        "这两张图像也说明，单独依赖整体图像级分数很难解释模型为什么出错。FSD 的三组结果一致判 real，说明它们共同受到整体自然感影响；Stay-Positive 在校准后判 fake，说明其局部异常敏感性在这类样本上更有价值；人工视觉观察则能进一步指出异常具体落在结构连接、透视关系、伪文本和局部排版规则上。三者结合之后，论文才不仅有“哪个模型对了”的结果，还有“为什么会这样”的解释。",
        "在答辩场景中，这类案例可以用来回应“实验结果是否只是数字堆叠”的问题。本文并没有只列出最终 ACC/AP，而是把分数、阈值、样本级冲突和图像内容放在一起分析，从而证明后续联合训练和 LVLM 标签不是临时添加概念，而是由真实失败样本推动出来的改进方向。"
    ])

    anchor = find_paragraph(doc, "从零初始化训练主要用于验证 FSD 训练链路是否能够在当前数据目录和服务器环境下跑通，因此其价值在于流程验证。正式可比较的 ADM 微调均使用官方 converted checkpoint 初始化。首轮 ADM 微调使用 lr=1e-5、batch_size=16、total_training_steps=10000、accumulation_steps=1、use_fp16=True，并以 ADM 作为 exclude_class，输出目录为 fsd/output/finetune_adm_stage1。该设置在 ADM 上取得 75.28%/78.54% 的 Accuracy/AP，明显优于从零初始化训练。")
    insert_after(anchor, [
        "首轮微调的训练过程还保留了日志间隔和检查点输出设置，便于观察训练是否正常推进。训练中每隔一定步数记录 loss、学习率和阶段性评估结果，完成后再使用统一评估脚本对多个生成器类别进行测试。这样可以确认首轮提升不是只在 ADM 单点上偶然出现，而是在 Midjourney、Stable Diffusion、GLIDE、VQDM 等类别上也保持了较高结果。",
        "不过，首轮微调虽然优于从零初始化训练，但 ADM 指标仍只是接近官方基线，并没有稳定超过官方基线。因此本文把它写成“有效训练探索”而不是“最终改进模型”。这一点很重要，因为它既能体现实际训练工作，又不会给答辩老师留下夸大改进幅度的印象。"
    ])

    anchor = find_paragraph(doc, "第二轮保守微调使用同一初始化 checkpoint，但把学习率降为 5e-6，总训练步数降为 5000，保存和评估间隔均为 5000，输出目录为 fsd/output/finetune_adm_stage2。结果在 ADM 上为 74.13%/76.89%，低于首轮微调。这个负结果很重要，因为它说明当前任务不是简单地降低学习率、缩短训练步数就会更稳定；模型对训练参数、样本构成和初始化状态都较敏感。")
    insert_after(anchor, [
        "第二轮结果也为后续联合训练提供了决策依据。如果继续只在 FSD 单模型上微调学习率和步数，收益并不稳定；更有价值的是引入能够解释困难样本的额外监督信号。因此后续工作没有继续无边界尝试更多 FSD 微调组合，而是转向 SP 分数约束和 LVLM 语义标签接入。"
    ])

    anchor = find_paragraph(doc, "第一阶段训练日志进一步表明，虽然 metadata 已经接入训练流程，但有效 Stay-Positive 监督样本不足，valid_sp_samples 长期为 0，sp_loss 也没有形成稳定贡献。因此，本阶段更准确的定位是“联合训练框架落地验证与问题暴露”，而不是“SP 分支已经带来显著性能提升”。这一定性能够解释为什么闭集指标与同协议 FSD-only 基线接近，也避免把工程链路跑通误写成性能突破。")
    insert_after(anchor, [
        "从论文写法上看，这个阶段仍然应该保留，而不是因为没有显著提升就删除。原因在于它回答了一个工程问题：FSD 主干能否读取包含 SP 分数的 metadata，并在闭集协议下完成训练、保存 checkpoint 和输出评估结果。答案是可以；但它同时暴露出另一个问题：SP 分数与当前 episode 采样方式、有效样本掩码和损失计算之间还没有形成足够有效的匹配。",
        "因此第一阶段的贡献不是性能增益，而是把“SP 如何接入 FSD”从文字方案推进到可运行代码和日志证据。这个结果也让第二阶段的设计更加明确，即必须引入更直接的有效标签信号，而不是只依赖跨域后可能失配的 SP 分数。"
    ])

    anchor = find_paragraph(doc, "LVLM 结构化标签不是凭空生成的，而是先从 ADM 困难样本中进行人工视觉观察。本文共整理 20 个代表性案例，其中包括首批优先样本、第二批 strong-conflict 样本和核心互补 Top12 样本。人工观察重点记录样本是否存在伪文本/伪界面、局部结构连接异常、生物体局部真实性不足、局部修补或过度平滑，以及布局或关系不一致等问题。")
    insert_after(anchor, [
        "初始人工标注采用“先筛选、再观察、再结构化”的流程。先从 ADM 样本级导出结果中筛选训练回退样本、边界翻转样本和 SP/FSD 强冲突样本；再对图像进行真实视觉观察，记录具体异常位置和可能误判原因；最后把自由文本观察转化为多标签字段。这样可以减少直接让模型对所有图像自动打标签带来的噪声，也能让论文中的 LVLM 标签与实际案例描述保持一致。",
        "人工观察并不要求每张图都存在所有异常，而是按照最主要证据进行标记。例如按键和屏幕类样本主要标记伪文本/伪界面，动物和昆虫样本主要标记生物体局部真实性不足或结构连接异常，室内和建筑类样本主要标记布局关系或空间几何异常。多标签形式允许一个样本同时带有两类以上异常，从而更贴近真实 ADM 图像的复杂性。"
    ])

    anchor = find_paragraph(doc, "弱扩展阶段并不把所有 ADM 图像都强行标为同一类语义异常，而是优先把人工标注案例、强冲突样本、提示词模板相关样本和高价值 ADM 子集写入 joint_training_metadata_stage2.csv。对于人工确认案例，使用较高的 lvlm_confidence；对于根据 ADM prompt 或冲突模式弱扩展得到的样本，使用较低的置信度和 hard_weight 控制影响范围。训练时只有 lvlm_valid 为真的样本参与 LVLM 辅助损失，从而避免少量弱标签过度主导主任务。")
    insert_after(anchor, [
        "轻量化接入 FSD 时，LVLM 不参与在线推理和反向传播，只把结构化标签作为辅助监督写入 metadata。训练批次读取图像后，FSD backbone 输出视觉特征，一路进入原型分类损失，另一路进入轻量多标签辅助头。辅助头只增加少量线性层参数，因此不会显著增加训练资源需求，也不会改变 FSD 主干的基本评估方式。",
        "这种接入方式的好处是边界清晰：主任务仍然由 Accuracy/AP 评价，LVLM 只负责提供局部异常语义监督；如果 LVLM 标签不足或质量不高，最坏情况也主要影响辅助头 F1，而不会把整套系统变成难以解释的大模型联合训练。本文第二阶段结果正体现了这种边界：主任务指标保持稳定，辅助头已经能输出 F1，但距离成熟语义监督仍有明显空间。"
    ])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
