"""对 4.29 改稿版论文做轻量版式优化。

本脚本不改变正文内容，只在摘要、目录和主要章节前增加分页，使渲染时不至于过密。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "AI伪造图像识别论文_终稿_改写2.docx"


def add_break_before(paragraph) -> None:
    """在指定段落前插入分页，若前一段已分页则不重复处理。"""

    previous = paragraph._p.getprevious()
    if previous is not None and "lastRenderedPageBreak" in previous.xml:
        return
    run = paragraph.insert_paragraph_before().add_run()
    run.add_break()


def main() -> None:
    """给摘要、目录和一级章节增加更清晰的分页。"""

    doc = Document(DOCX_PATH)
    break_titles = {
        "面向泛化检测的AI伪造图像多模型识别研究",
        "目  录",
        "绪论",
        "基础知识与相关工作",
        "面向泛化检测的多模型协同分析方案",
        "实验设置与设计",
        "实验结果与分析",
        "总结与展望",
        "参考文献",
    }

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in break_titles:
            add_break_before(paragraph)

    doc.save(DOCX_PATH)
    print("layout_polished=1")


if __name__ == "__main__":
    main()
