"""将联合训练第二阶段结果回填到 Word 终稿。

该脚本只面向本项目当前的论文终稿使用，核心目标是：
1. 在不重建整篇 Word 的前提下，替换已经过时的正文表述；
2. 补入两阶段联合训练，尤其是第二阶段 LVLM 最小量化训练结果；
3. 生成一张联合训练对比图，并在第 5 章中插入结果表与图片；
4. 保留原始文档备份，便于需要时人工回退或对照。

注意：脚本不会删除任何文件，只会创建备份、图片和更新目标 docx。
"""

from __future__ import annotations

import shutil
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "AI伪造图像识别论文_终稿.docx"
BACKUP_PATH = ROOT / "docs" / "AI伪造图像识别论文_终稿_联合训练回填前备份.docx"
FIGURE_PATH = ROOT / "figures" / "joint_stage2_min_comparison.png"


def count_chinese_chars(doc: Document) -> int:
    """统计正文、表格中的中文字符数量，用于粗略控制论文篇幅。"""

    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return sum(1 for char in "\n".join(chunks) if "\u4e00" <= char <= "\u9fff")


def find_paragraph(doc: Document, text: str) -> int:
    """按照完整段落文本定位段落，避免依赖易变化的段落序号。"""

    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return index
    raise ValueError(f"未找到段落：{text}")


def find_paragraph_contains(doc: Document, needle: str, start: int = 0) -> int:
    """按照关键短语定位段落，适合查找正文中的章节锚点。"""

    for index, paragraph in enumerate(doc.paragraphs[start:], start=start):
        if needle in paragraph.text:
            return index
    raise ValueError(f"未找到包含关键字的段落：{needle}")


def replace_paragraph_text(paragraph, text: str) -> None:
    """替换段落文字，同时尽量保留该段落原有样式、缩进和对齐方式。"""

    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    apply_run_font(run)


def apply_run_font(run, size: Pt | None = None, bold: bool | None = None) -> None:
    """统一新写入文字的中英文字体，减少新增内容与原文格式割裂。"""

    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def paragraph_after(paragraph, text: str = "", style: str | None = None):
    """在指定段落后插入新段落，并返回新段落对象。"""

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
        for run in new_paragraph.runs:
            apply_run_font(run)
    return new_paragraph


def paragraph_after_element(element, parent, text: str = "", style: str | None = None):
    """在任意块级 XML 元素后插入段落，主要用于表格后追加题注。"""

    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    new_paragraph = Paragraph(new_p, parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
        for run in new_paragraph.runs:
            apply_run_font(run)
    return new_paragraph


def insert_table_after(paragraph, rows: list[list[str]], caption: str):
    """在指定段落后插入表格和题注，并设置为论文中较稳妥的紧凑样式。"""

    table = paragraph._parent.add_table(rows=len(rows), cols=len(rows[0]), width=Cm(15.5))
    # 不同 Word 模板中的表格样式名称可能不同；这里优先沿用当前论文已有表格的样式，
    # 若模板不包含该样式，则保留默认样式，避免因为样式名差异中断终稿生成。
    for style_name in ("Normal Table", "Table Grid"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for cell_paragraph in cell.paragraphs:
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell_paragraph.runs:
                    apply_run_font(run, Pt(9.5), bold=(row_index == 0))
            if row_index == 0:
                shade_cell(cell, "D9EAF7")

    # 将表格移动到锚点段落之后；python-docx 默认追加到文末，因此需要调整 OOXML 位置。
    paragraph._p.addnext(table._tbl)

    caption_paragraph = paragraph_after_element(table._tbl, paragraph._parent, caption, "Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table, caption_paragraph


def shade_cell(cell, fill: str) -> None:
    """设置表头底色，让新增表格与论文现有表格区分但不过度花哨。"""

    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def insert_picture_after(paragraph, image_path: Path, caption: str, width_cm: float = 13.2):
    """在指定段落后插入图片和图题，图片尺寸控制在正文宽度内。"""

    image_paragraph = paragraph_after(paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    caption_paragraph = paragraph_after(image_paragraph, caption, "Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return image_paragraph, caption_paragraph


def remove_paragraph(paragraph) -> None:
    """删除脚本自己定位到的旧段落，避免同一位置同时保留新旧口径。"""

    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def make_joint_stage2_figure() -> None:
    """生成联合训练第二阶段结果图，用于在论文中直观展示主任务指标走势。"""

    FIGURE_PATH.parent.mkdir(parents=True, exist_ok=True)

    generators = ["ADM", "SD", "Midjourney"]
    stage1_acc = [95.50, 95.34, 87.00]
    stage2_acc = [95.22, 95.33, 86.97]
    stage1_ap = [97.50, 97.90, 90.47]
    stage2_ap = [97.45, 97.72, 90.38]

    plt.rcParams["font.family"] = "DejaVu Sans"
    plt.rcParams["axes.unicode_minus"] = False

    fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.2), dpi=180)
    x_positions = range(len(generators))
    bar_width = 0.34

    for axis, title, baseline, stage2, y_label in [
        (axes[0], "Accuracy", stage1_acc, stage2_acc, "Percent (%)"),
        (axes[1], "Average Precision", stage1_ap, stage2_ap, "Percent (%)"),
    ]:
        axis.bar(
            [x - bar_width / 2 for x in x_positions],
            baseline,
            width=bar_width,
            label="Stage1 FSD+SP",
            color="#4B7BAF",
        )
        axis.bar(
            [x + bar_width / 2 for x in x_positions],
            stage2,
            width=bar_width,
            label="Stage2 +LVLM",
            color="#D97842",
        )
        axis.set_title(title, fontsize=12, fontweight="bold")
        axis.set_ylabel(y_label, fontsize=10)
        axis.set_xticks(list(x_positions), generators)
        axis.set_ylim(80, 100)
        axis.grid(axis="y", alpha=0.25, linestyle="--", linewidth=0.7)
        axis.spines["top"].set_visible(False)
        axis.spines["right"].set_visible(False)
        for x, value in zip(x_positions, baseline):
            axis.text(x - bar_width / 2, value + 0.35, f"{value:.2f}", ha="center", fontsize=8)
        for x, value in zip(x_positions, stage2):
            axis.text(x + bar_width / 2, value + 0.35, f"{value:.2f}", ha="center", fontsize=8)

    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=2, frameon=False, fontsize=10)
    fig.suptitle("Closed-set Results of Two-stage Joint Training", fontsize=13, fontweight="bold")
    fig.tight_layout(rect=[0, 0.09, 1, 0.92])
    fig.savefig(FIGURE_PATH, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def update_toc_text(doc: Document) -> None:
    """同步手工目录中的小节名称，让目录与新增正文口径保持一致。"""

    replacements = {
        "5.5 实验结果的研究启示": "5.5 联合训练结果与研究启示",
        "5.6 实验结果的局限性分析": "5.6 实验结果的局限性分析",
        "6 改进思路与研究计划": "6 改进思路与后续展望",
    }
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if stripped in replacements:
            replace_paragraph_text(paragraph, paragraph.text.replace(stripped, replacements[stripped]))


def update_core_paragraphs(doc: Document) -> None:
    """替换摘要、方法定位、实验边界、局限性和结论中的旧口径。"""

    replacements = {
        15: (
            "随着扩散模型和多模态生成技术的发展，AI 伪造图像检测面临更强的跨模型泛化、判别依据稳定性和复杂样本解释挑战。本文基于公开代码、预训练权重和 GenImage 数据，对 FSD、Stay-Positive 与 LVLM 三条路线开展基线复现、样本级分析与联合训练验证。结果表明，FSD 在 Midjourney、Stable Diffusion、ADM、BigGAN、GLIDE 和 VQDM 上分别获得 79.56%、88.34%、75.41%、79.27%、96.67% 和 75.47% 的准确率；Stay-Positive 在固定真假对照任务中表现更强，但在 ADM 上出现明显阈值漂移，默认阈值下整体准确率为 52.25%，经校准后可提升到 80.77%。在此基础上，本文进一步完成两阶段联合训练验证：第一阶段 FSD + Stay-Positive 在闭集协议下完成 10000 step 训练，在 ADM、SD、Midjourney 上分别取得 95.50%/97.50%、95.34%/97.90%、87.00%/90.47% 的 Accuracy/AP；第二阶段在 FSD 主干上引入 LVLM 结构化语义标签与轻量辅助头，训练日志显示 steps_with_valid_lvlm=6094，ADM 上输出 LVLM F1=0.1778，并在三类闭集评估中取得 95.22%/97.45%、95.33%/97.72%、86.97%/90.38% 的 Accuracy/AP。实验说明，LVLM 语义监督已经真实进入训练计算图，但主任务性能尚未显著超过同协议基线。因此，本文最终形成了从基线复现、ADM 冲突分析到两阶段联合训练的完整研究闭环。"
        ),
        19: (
            "AI-generated images are increasingly difficult to detect across unseen generators. This thesis reproduces and compares FSD, Stay-Positive, and LVLM on GenImage using public code, pretrained checkpoints, sample-level analysis, and two-stage joint training. FSD reaches 79.56%-96.67% accuracy across six generators, while Stay-Positive performs strongly in paired binary settings but shows clear threshold drift on ADM. The first joint stage, FSD + Stay-Positive, completes 10000 training steps and obtains 95.50%/97.50%, 95.34%/97.90%, and 87.00%/90.47% Accuracy/AP on ADM, SD, and Midjourney. The second minimal quantitative stage introduces structured LVLM labels and a lightweight auxiliary head. The training log reports steps_with_valid_lvlm=6094, and the final model outputs LVLM F1=0.1778 on ADM while keeping the main-task metrics stable. These results indicate that LVLM supervision has entered the training graph, although it has not yet produced a significant main-task improvement. The thesis therefore builds a complete pipeline from baseline reproduction and ADM conflict analysis to joint-training validation."
        ),
        83: (
            "本文的目标不是简单堆叠新的网络结构，而是在公开资源条件下完成 FSD、Stay-Positive 和 LVLM 的代表性复现与联合验证，明确不同方法在泛化检测、鲁棒判别和语义解释上的能力边界，并进一步检验三类信息能否在同一训练框架中形成可运行的闭环。"
        ),
        85: (
            "围绕这一目标，本文主要完成六项工作：梳理 AI 伪造图像检测相关研究；搭建 GenImage 最小可运行实验环境；复现 FSD 与 Stay-Positive 代表性结果；整理 ADM 场景下的校准、分数分布与样本级冲突现象；构造 LVLM 结构化案例分析与弱标签扩展；最后完成两阶段联合训练实验，验证 Stay-Positive 分数约束和 LVLM 语义标签接入 FSD 主干后的实际训练信号与闭集表现。具体而言，本文既关注方法层面的定位差异，也关注工程层面的可运行性、结果可复核性与样本级证据是否能够支撑结论，使论文不只是停留在结果罗列上，而是形成从“能跑通”“能比较”“能解释”到“能联合训练”的完整链条。"
        ),
        86: (
            "研究方法上，本文采用文献分析、代码复现、结果比较、样本级统计和最小量化训练相结合的方式推进。对于每条技术路线，均先确认其公开资源是否可运行，再根据已有结果判断其在本项目中的合理定位；在联合训练阶段，则优先采用风险较低、证据更明确的两阶段路线：先验证 FSD 与 Stay-Positive 的闭集联合训练，再把 LVLM 结构化标签接入为轻量辅助监督。这样既避免在时间有限条件下盲目扩大模型规模，也保证最终论文中的结论能够由日志、结果表和样本级证据共同支撑。"
        ),
        88: (
            "本文的特点在于：先做实基线，再讨论改进；以公开代码、公开权重和公开数据作为实验基础；在结果解释上保持保守，不把未显著提升的结果包装成突破；同时又进一步完成了两阶段联合训练，使论文不再只停留在基线复刻层面。尤其是第二阶段最小量化版实验表明，LVLM 语义标签已经能够作为辅助监督进入训练计算图，这为后续从纯视觉判别走向视觉-语义联合检测提供了可复核的起点。"
        ),
        90: (
            "全文依次讨论研究背景与意义、相关工作、方法定位、实验设计、结果分析、联合训练验证、改进展望以及结论。"
        ),
        111: (
            "这类方法更适合处理边界样本、冲突样本和难以用单一分数解释的复杂案例。对于论文写作而言，LVLM 的价值也在于帮助把“模型为什么可能错”转化为更可描述、更贴近人工观察语言的分析文本。本文在前期把 LVLM 用于结构化案例观察，在后期进一步将其标签扩展为可训练的辅助监督，因此 LVLM 不再只承担后验解释角色，而是成为联合训练框架中的语义信号来源。"
        ),
        112: (
            "但 LVLM 的输出仍受提示词、视觉上下文和标签覆盖范围影响，当前更适合作为轻量辅助监督与解释接口，而非直接替代专门训练的检测器。因此，本文对 LVLM 的定位保持克制：强调它已经完成最小量化接入，但不夸大为主任务性能显著提升。"
        ),
        130: (
            "LVLM 被定位为“复杂场景辅助分析与轻量语义监督接口”，用于补充传统判别模型难以说明的边界样本和冲突样本，并在第二阶段联合训练中提供结构化语义标签。"
        ),
        131: (
            "因此，LVLM 在本文中的主要价值不是与 FSD、Stay-Positive 做机械化数值排名，而是帮助模型从纯视觉分数进一步接触局部结构异常、伪文本、局部修补和生物体细节失真等更接近人工观察的语义线索。"
        ),
        133: (
            "三类方法的关系可以概括为：FSD 提供开放场景主干，Stay-Positive 提供鲁棒判别思路与离线分数参考，LVLM 提供复杂案例解释能力和轻量语义监督。"
        ),
        134: (
            "本文最终采用两阶段方式探索三者结合：第一阶段验证 FSD 与 Stay-Positive 的闭集联合训练，第二阶段在 FSD 主干中接入 LVLM 结构化标签，以最小风险完成从案例解释到可训练语义监督的过渡。"
        ),
        136: (
            "本文的研究边界是：已经完成多模型基线复现、ADM 样本级冲突分析、第一阶段 FSD + Stay-Positive 联合训练，以及第二阶段 FSD + LVLM 最小量化训练；但当前仍不宣称 LVLM 已经显著提升主任务性能，也不把不同协议下的结果解释为绝对排名。论文重点在于给出真实、可复核、边界清楚的阶段性结论。"
        ),
        138: (
            "整体方案可概括为四步：完成基线复现；在统一数据上比较方法差异；结合样本级分析定位 ADM 失败模式；最后以两阶段联合训练验证 Stay-Positive 分数约束和 LVLM 语义监督接入 FSD 主干后的可行性。"
        ),
        152: (
            "FSD 的实验流程包括数据整理、checkpoint 适配、正式评估、ADM 微调探索和联合训练验证。具体执行时，先完成公开权重的兼容转换，再验证官方基线是否能够在当前环境中稳定跑通，随后围绕 ADM 构造额外训练与评估链路，用于观察该方法在困难类别上的可调性；最后通过两阶段联合训练，把 Stay-Positive 离线分数和 LVLM 结构化标签逐步接入 FSD 主干。整个流程的核心目标不是盲目追求更高数值，而是确保每一步结果都可解释、可复核。"
        ),
        153: (
            "Stay-Positive 的实验流程包括生成 CSV、运行预训练模型、汇总分数并补充 ADM 校准分析。由于其真实样本与伪造样本采用对照式输入，因此在结果整理阶段还需要特别关注分数分布、阈值设置和真假样本比例之间的关系。联合训练阶段进一步把 Stay-Positive 分数作为离线参考信号纳入元数据，但日志显示该信号在当前配置下没有形成有效的非零监督，这也成为第一阶段结果解释中的重要边界。"
        ),
        155: (
            "评价指标主要采用 Accuracy 与 AP，前者反映固定阈值下的整体分类正确率，后者反映不同阈值下的综合排序能力。对于第二阶段 LVLM 最小量化训练，本文额外记录 LVLM 辅助头的多标签 F1，用于判断语义标签是否真实进入训练并产生可度量输出。"
        ),
        161: (
            "对 LVLM 而言，本文已完成从案例分析到最小量化训练的接入验证；其辅助头能够输出 ADM 上的 LVLM F1，但由于标签规模有限，当前更适合用于证明语义监督可接入，而不是证明主任务性能已经显著提升。"
        ),
        186: "联合训练结果与研究启示",
        201: (
            "尽管本文已经获得较为明确的结果，但仍存在局限：不同方法的原始覆盖范围与测试协议并不完全一致，ADM 相关证据也主要集中在单一高价值场景。此外，联合训练虽然补齐了“是否能够接入”的关键证据，但仍属于最小量化验证，而不是大规模、多数据集、充分调参后的最终性能上限。因此在写作中必须保持结论边界，既说明三类模型已经形成可运行闭环，也不把阶段性结果夸大为绝对领先。"
        ),
        204: (
            "目前，这种更深层差异主要由案例观察、局部统计和最小量化训练共同支持，后续仍需要更大规模标签、更稳定的提示词协议和更多生成器类别来验证。"
        ),
        205: (
            "LVLM 已经形成可训练的最小量化结果，但标签主要来自结构化案例与弱标签扩展，覆盖面仍有限；ADM 上 LVLM F1=0.1778 说明辅助头已经开始学习语义信号，但距离稳定成熟的视觉-语义检测器仍有明显差距。"
        ),
        206: (
            "总体上，本文在结果解释上坚持“已完成什么就写什么”的原则，即承认 Stay-Positive 的高性能，也保留 FSD 的研究价值，并把 LVLM 的贡献定位为从解释接口走向轻量辅助监督的初步验证。这种处理方式虽然相对克制，但更符合当前证据水平，也能避免论文结论与实际完成度之间出现脱节。对于毕业论文来说，真实、稳健、边界清楚往往比盲目追求更激进的结论更重要。"
        ),
        208: (
            "截至目前，本文已经完成多模型框架重构、最小数据子集整理、正式基线复现、ADM 校准分析、样本级冲突统计、训练探索以及两阶段联合训练验证，为最终定稿提供了较完整的实验与文本支撑。"
        ),
        209: "改进思路与后续展望",
        211: (
            "在完成最终论文之前，仍需清楚认识几个现实问题：实验覆盖范围仍有限，不同方法协议并不完全一致，第二阶段 LVLM 量化训练也仍处于最小验证级别。这些问题并不否定当前工作的价值，反而说明本文最重要的成果不是“某个指标绝对最高”，而是把三条代表性技术路线组织成了一个可运行、可比较、可解释、可继续扩展的研究框架。换言之，当前阶段最重要的不是继续无边界地增加实验数量，而是把已有结果组织成更有解释力、更能经得起追问的最终论文结论。"
        ),
        212: (
            "因此，后续若继续推进，更适合围绕已有闭环做质量提升，而不是重新铺开高风险的大规模训练。"
        ),
        213: (
            "已有训练探索说明，FSD 主干训练链路、闭集联合训练链路和 LVLM 辅助头训练链路均已打通；后续改进应优先提升监督质量、标签覆盖和损失权重设计。"
        ),
        215: (
            "基于现有结果，后续总体思路是继续以 FSD 为主干，保留 Stay-Positive 的鲁棒判别思想，并把 LVLM 结构化标签扩展为更稳定的语义辅助监督。具体而言，可以从三个方向深化：一是扩大 LVLM 标签覆盖范围，让语义监督不只集中在 ADM 代表样本；二是改进 Stay-Positive 分数对齐方式，避免第一阶段出现有效监督样本不足的问题；三是围绕 ADM 这类高价值困难场景继续做误判分析、局部异常定位和阈值校准。这样既能控制风险，也能让研究持续沿着最有价值的问题推进。"
        ),
        216: (
            "具体工作包括：继续梳理跨类别差异；围绕误判、分数分布和阈值敏感性做轻量分析；把 ADM 失败模式与案例观察进一步压缩进论文讨论；在条件允许时扩大 LVLM 标签数量，并尝试更细粒度的语义损失设计。"
        ),
        218: (
            "结合项目进度，下一阶段若继续开展实验，不宜盲目扩大模型规模，而应优先完善第二阶段量化验证：增加 LVLM 标签样本、检查辅助头类别不均衡问题、补充不同权重下的消融实验，并观察主任务指标与语义 F1 是否能够同时改善。"
        ),
        220: (
            "从论文写作角度看，当前最终版应围绕已完成结果进行收束：基线复现用于说明三类方法的初始能力，ADM 样本级分析用于解释方法差异，两阶段联合训练用于回答“是否真正接入联合框架”。这种顺序既能锁定已经完成且可复核的成果，也能让读者看到本文从复现比较走向改进验证的研究推进过程。"
        ),
        222: (
            "预期成果是形成一份结构完整、结果真实、边界清楚的毕业论文，并为后续围绕 FSD 主干、Stay-Positive 分数校准和 LVLM 语义监督的进一步改进保留明确入口。"
        ),
        224: (
            "本文围绕未知生成器泛化、鲁棒判别与复杂样本解释三个方向，对 FSD、Stay-Positive 与 LVLM 开展了基线复现、样本级分析和联合训练验证。论文重点并非宣称提出一套完全成熟的新模型，而是在公开资源条件下把关键基线跑通、把结果讲清楚、把方法边界说明白，并进一步完成两阶段联合训练，使研究从单纯复刻比较推进到可运行的融合验证。这种处理方式虽然更克制，却更能体现阶段性研究工作的真实价值。"
        ),
        225: (
            "结果表明，FSD 在六类生成器上表现出明显分层，Stay-Positive 在固定真假对照任务中更强，而 ADM 场景则同时暴露出校准问题与稳定的样本级互补性。尤其是 Rajan/Ours+ 在 ADM 上默认阈值下仅得 52.25%，经 0.388818 校准后可提升到 80.77%；同时，3000 个 ADM 假样本中存在 1137 个“Stay-Positive 判 fake 而 FSD 三组均判 real”的主导互补样本。训练探索还说明，基于官方 checkpoint 的微调明显优于随机初始化，但更保守的第二轮微调并不会自然带来提升。进一步的联合训练结果表明，第一阶段 FSD + Stay-Positive 在 ADM、SD、Midjourney 上取得 95.50%/97.50%、95.34%/97.90%、87.00%/90.47% 的 Accuracy/AP；第二阶段 FSD + LVLM 最小量化训练取得 95.22%/97.45%、95.33%/97.72%、86.97%/90.38%，并输出 ADM 上的 LVLM F1=0.1778。由此可见，联合训练链路已经完整跑通，LVLM 语义监督也已经进入训练计算图，但主任务性能仍主要保持稳定，尚未形成显著超越同协议基线的结论。"
        ),
        226: (
            "综合来看，Stay-Positive 更适合作为高性能判别基线，FSD 更适合作为泛化检测研究主干，LVLM 则从复杂样本的语义分析补充进一步扩展为轻量辅助监督来源。本文已经完成了从基线复现到样本级分析，再到两阶段联合训练验证的关键过渡。对于毕业论文而言，这样的组织方式较为稳妥：既保留真实完成的实验成果，又避免夸大尚未充分验证的性能收益；既说明三类方法各自的能力边界，也为未来继续深化 FSD 主干、改进 Stay-Positive 分数对齐和扩展 LVLM 语义标签留下清晰入口。后续若有更多时间和资源，围绕 ADM 的误判模式、LVLM 标签扩展和多损失权重消融继续做有控制的深化，将是最具性价比的研究方向。"
        ),
    }

    for index, text in replacements.items():
        replace_paragraph_text(doc.paragraphs[index], text)


def insert_joint_training_section(doc: Document) -> None:
    """在第 5 章中插入联合训练结果表、图和解释段落。"""

    anchor_index = find_paragraph_contains(doc, "因此，训练探索部分最重要的结论")
    # 在 FSD 训练探索结论之后插入新小节，使叙事顺序保持为：
    # 基线结果 -> 微调探索 -> 联合训练 -> 研究启示与局限性。
    heading_anchor = doc.paragraphs[anchor_index]

    p_heading = paragraph_after(heading_anchor, "两阶段联合训练结果", "Heading 2")
    p_intro = paragraph_after(
        p_heading,
        (
            "在完成基线复现、ADM 校准分析和微调探索之后，本文进一步把三类方法推进到联合训练层面。联合训练并不是简单把三个模型结果相加，而是尝试回答一个更关键的问题：FSD 的开放场景主干、Stay-Positive 的鲁棒判别思想以及 LVLM 的语义解释信号，能否在同一训练流程中形成可运行、可记录、可评价的闭环。为降低风险，本文采用两阶段策略。第一阶段先引入 Stay-Positive 离线分数，验证闭集联合训练流程是否稳定；第二阶段再引入 LVLM 结构化标签和轻量辅助头，检验语义监督是否能够真实进入训练计算图。"
        ),
        "Normal",
    )

    p_stage1 = paragraph_after(
        p_intro,
        (
            "第一阶段 FSD + Stay-Positive 训练完成 10000 step，并在 ADM、Stable Diffusion、Midjourney 三类闭集评估上分别取得 95.50%/97.50%、95.34%/97.90%、87.00%/90.47% 的 Accuracy/AP。从主任务指标看，该结果与同协议 FSD-only 基线基本一致，说明第一阶段主要验证了训练流程可运行，而不是证明 Stay-Positive 分数已经带来明显增益。结合日志中 SP 有效样本不足的现象，可以判断当前离线分数约束仍需要进一步做阈值、覆盖范围和采样策略上的适配。"
        ),
        "Normal",
    )
    p_stage2 = paragraph_after(
        p_stage1,
        (
            "第二阶段最小量化版进一步引入 LVLM 结构化语义标签。服务器作业 17862387 完成 10000 step 训练，日志显示训练过程中多次出现非零 valid_lvlm_samples 和 lvlm_loss，最终 steps_with_valid_lvlm 达到 6094，avg_valid_lvlm_samples_per_step 为 1.6343。这说明 LVLM 辅助头并非停留在论文方案或后验解释层面，而是已经参与反向传播。最终 10000 step 评估中，ADM、SD、Midjourney 的 Accuracy/AP 分别为 95.22%/97.45%、95.33%/97.72%、86.97%/90.38%，ADM 上额外输出 LVLM F1=0.1778。"
        ),
        "Normal",
    )

    rows = [
        ["训练阶段", "ADM Acc/AP", "SD Acc/AP", "Midjourney Acc/AP", "LVLM F1", "结论"],
        ["第一阶段：FSD + Stay-Positive", "95.50% / 97.50%", "95.34% / 97.90%", "87.00% / 90.47%", "N/A", "闭集训练稳定，主任务指标与同协议基线基本一致"],
        ["第二阶段：FSD + LVLM 最小量化", "95.22% / 97.45%", "95.33% / 97.72%", "86.97% / 90.38%", "0.1778", "语义辅助头真实进入训练，但主任务尚未显著提升"],
    ]
    table, table_caption = insert_table_after(p_stage2, rows, "表 5-4 两阶段联合训练结果对比")

    p_after_table = paragraph_after(
        table_caption,
        (
            "从表 5-4 可以看出，第二阶段主任务指标相较第一阶段略有波动，但总体仍处于同一水平区间。这个结果的解释应当保持谨慎：它不能被写成“LVLM 显著提升了伪造图像检测性能”，但可以支撑另一个同样重要的结论，即 LVLM 语义标签已经完成从结构化案例分析到可训练辅助监督的关键过渡。对于本文而言，这一结论具有实际意义，因为它使论文不再只是比较三篇已有工作的公开结果，而是进一步完成了面向三模型融合的最小实现和量化验证。"
        ),
        "Normal",
    )

    image_paragraph, image_caption = insert_picture_after(
        p_after_table,
        FIGURE_PATH,
        "图 5-4 两阶段联合训练闭集指标对比图",
        width_cm=13.0,
    )

    paragraph_after(
        image_caption,
        (
            "图 5-4 进一步展示了两阶段结果的变化趋势。可以看到，加入 LVLM 后各生成器上的 Accuracy 与 AP 没有出现剧烈下降，说明辅助头在当前配置下没有破坏 FSD 主干的基本判别能力；但指标也没有形成稳定提升，说明语义监督仍受标签规模、样本覆盖和损失权重限制。因而本文把第二阶段定位为“最小量化验证”而不是“最终性能优化”。这种定位更符合实验数据本身，也为后续继续扩充 LVLM 标签、设计更细粒度语义损失和进行消融实验提供了清晰方向。"
        ),
        "Normal",
    )


def clean_duplicate_or_stale_phrases(doc: Document) -> None:
    """对残留旧口径做最后兜底替换，避免终稿出现前后矛盾。"""

    stale_replacements = {
        "当前不宣称完成联合训练，也不宣称已经得到 LVLM 的系统性定量结果": "已经完成两阶段联合训练和 LVLM 最小量化验证，但仍不宣称主任务性能显著提升",
        "LVLM 当前也尚未形成与前两类模型可直接对比的统一定量结果": "LVLM 已经形成最小量化结果，但仍不适合与前两类模型做简单绝对排名",
        "LVLM 仍缺少统一定量结果": "LVLM 已完成最小量化结果，但仍需要扩大标签规模",
        "后续若条件允许，可以进一步探索": "本文已初步探索",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in stale_replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            replace_paragraph_text(paragraph, new_text)


def main() -> None:
    """执行 Word 回填主流程。"""

    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    make_joint_stage2_figure()

    doc = Document(DOCX_PATH)
    before_chars = count_chinese_chars(doc)

    update_toc_text(doc)
    update_core_paragraphs(doc)
    insert_joint_training_section(doc)
    clean_duplicate_or_stale_phrases(doc)

    doc.save(DOCX_PATH)

    updated = Document(DOCX_PATH)
    after_chars = count_chinese_chars(updated)
    full_text = "\n".join(paragraph.text for paragraph in updated.paragraphs)
    stale_hits = [
        phrase
        for phrase in [
            "当前不宣称完成联合训练",
            "尚未形成统一口径的正式数值评测",
            "LVLM 当前也尚未形成",
            "LVLM 仍缺少统一定量结果",
        ]
        if phrase in full_text
    ]

    print(f"backup={BACKUP_PATH}")
    print(f"figure={FIGURE_PATH}")
    print(f"before_cn_chars={before_chars}")
    print(f"after_cn_chars={after_chars}")
    print(f"tables={len(updated.tables)}")
    print(f"inline_shapes={len(updated.inline_shapes)}")
    print(f"stale_hits={stale_hits}")


if __name__ == "__main__":
    main()
