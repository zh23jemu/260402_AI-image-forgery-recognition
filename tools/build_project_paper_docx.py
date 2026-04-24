#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
基于论文模板副本生成项目论文成稿。

这个脚本的目标不是做一个通用的 Markdown 转 DOCX 引擎，而是稳定地把
当前项目已经整理好的论文主稿、实验表格和图片落到模板副本中，生成一个
可以继续人工微调的 Word 成品。
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE_CN = "基于多模型基线复现与改进的AI伪造图像识别研究"
TITLE_EN = "Research on AI-Generated Image Detection Based on Multi-Model Baseline Reproduction and Improvement"
HEADER_TEXT = "AI伪造图像识别研究"
UNIVERSITY_NAME = "本科毕业论文"
COLLEGE_NAME = "人工智能与计算机相关专业方向"
AUTHOR_NAME = "学生姓名：__________"
STUDENT_ID = "学    号：__________"
SUPERVISOR_NAME = "指导教师：__________"
SUBMIT_DATE = "完成时间：2026 年 4 月"
REFERENCE_HEADING = "参考文献"

ABSTRACT_CN = (
    "随着扩散模型、多模态生成模型和高质量图像生成技术的快速发展，AI伪造图像在视觉真实感、"
    "语义一致性和细节表达等方面不断提升，传统依赖纹理伪迹、压缩噪声或固定生成器特征的检测方法"
    "面临明显的泛化挑战。针对该问题，本文以公开论文提供的官方代码、预训练权重和公开数据为基础，"
    "围绕FSD、Stay-Positive与LVLM三条技术路线开展基线复现、结果比较和方案设计。本文首先在"
    "GenImage数据上复现FSD的六类生成器检测结果，随后验证Stay-Positive在Midjourney、"
    "Stable Diffusion场景中的高精度表现，并围绕ADM数据补充阈值校准、分数分布和样本级冲突统计。"
    "结果表明，FSD在Midjourney、Stable Diffusion、ADM、BigGAN、GLIDE和VQDM上的准确率"
    "分别为79.56%、88.34%、75.41%、79.27%、96.67%和75.47%；Stay-Positive中Rajan/Ours+"
    "在ADM上默认阈值0.5下准确率仅为52.25%，但经阈值0.388818校准后可提升到80.77%。进一步统计"
    "发现，在3000个ADM假样本中存在1137个“Stay-Positive校准后判fake，而FSD三组结果均判real”"
    "的主导互补样本，说明两类方法在局部异常证据与整体自然感之间存在稳定取证差异。三轮真实案例"
    "观察还表明，这些冲突主要集中在伪文本/伪界面、局部结构连接、生物体局部真实性不足以及局部修补"
    "或过度平滑等细节层面。综合分析可知，Stay-Positive更适合作为高性能判别基线，FSD更适合作为"
    "未知生成器泛化研究主干，而LVLM更适合作为复杂样本语义分析与解释接口。"
)

ABSTRACT_EN = (
    "With the rapid development of diffusion models, multimodal generation systems, and high-quality image "
    "synthesis techniques, AI-generated fake images have become increasingly realistic in semantics, texture, "
    "and local details, which makes generalizable detection much harder. To address this challenge, this thesis "
    "reproduces and analyzes three representative technical routes, namely FSD, Stay-Positive, and LVLM, based "
    "on publicly released code, pretrained checkpoints, and the GenImage benchmark. We first reproduce the FSD "
    "baseline on six generators, including Midjourney, Stable Diffusion, ADM, BigGAN, GLIDE, and VQDM. We then "
    "verify the strong closed-set discrimination performance of Stay-Positive on Midjourney and Stable Diffusion, "
    "and further extend the analysis to ADM through threshold calibration, score-distribution inspection, and "
    "sample-level conflict statistics. Experimental results show that FSD achieves accuracies of 79.56%, 88.34%, "
    "75.41%, 79.27%, 96.67%, and 75.47% on the six evaluated generators respectively. For Stay-Positive, the "
    "Rajan/Ours+ model reaches only 52.25% accuracy on ADM under the default threshold of 0.5, while its accuracy "
    "can be improved to 80.77% after recalibrating the threshold to 0.388818. Moreover, among 3000 ADM fake "
    "samples, 1137 dominant complementary cases are observed, where calibrated Stay-Positive predicts fake but all "
    "three FSD variants predict real. Visual case studies further indicate that these conflicts are mainly related "
    "to pseudo text or pseudo interfaces, abnormal local connections, weak biological realism, and local repair or "
    "over-smoothing artifacts. Overall, Stay-Positive is more suitable as a high-performance discrimination baseline, "
    "FSD is more valuable as the main research backbone for unknown-generator generalization, and LVLM is more "
    "appropriate as an auxiliary interface for semantic analysis and explanation of difficult cases."
)

KEYWORDS_CN = ["AI伪造图像识别", "FSD", "Stay-Positive", "LVLM", "基线复现", "泛化检测"]
KEYWORDS_EN = [
    "AI-generated image detection",
    "FSD",
    "Stay-Positive",
    "LVLM",
    "generalization analysis",
]

ACKNOWLEDGEMENTS = (
    "在本课题推进过程中，指导老师在研究方向选择、实验路线梳理、论文结构组织和阶段性问题分析等方面"
    "给予了持续而细致的指导，使本文能够在有限时间内围绕“先做实基线、再讨论改进”的思路稳步推进。"
    "同时，项目相关公开论文、开源代码、预训练权重与公开数据集为本文实验复现提供了重要基础；服务器"
    "环境、虚拟环境配置、日志排查与结果整理工作也为论文写作提供了稳定支撑。最后，感谢在项目讨论、"
    "实验推进和文本校对过程中提供帮助的老师与同学。"
)

STATEMENT_TITLE = "诚 信 声 明"
STATEMENT_BODY = (
    "我声明，所呈交的毕业论文是本人在老师指导下进行的研究工作及取得的研究成果。"
    "据我查证，除了文中特别加以标注和致谢的地方外，论文中不包含其他人已经发表或撰写过的研究成果，"
    "也不包含为获得其他教育机构的学位或证书而使用过的材料。我承诺，论文中的所有内容均真实、可信。"
)
STATEMENT_SIGN = "毕业论文作者签名：                              签名日期：    年    月    日"


def remove_all_body_content(doc: Document) -> None:
    """清空模板正文节点，但保留模板自带的页面设置和节属性。"""

    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size_pt: int = 12, bold: bool = False, italic: bool = False) -> None:
    """统一设置中英文混排字体，避免中文在不同机器上退回到不可控字体。"""

    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def configure_styles(doc: Document) -> None:
    """在模板样式基础上做轻量配置，使正文、标题、题注和参考文献风格一致。"""

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(24)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for style_name, size_pt in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size_pt)
        style.font.bold = True
        pf = style.paragraph_format
        pf.line_spacing = 1.25
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)

    if "Caption" in doc.styles:
        style = doc.styles["Caption"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(11)
        style.font.bold = False
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.first_line_indent = Pt(0)

    if "参考文献" in doc.styles:
        style = doc.styles["参考文献"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.paragraph_format.first_line_indent = Pt(-21)
        style.paragraph_format.left_indent = Pt(21)
        style.paragraph_format.line_spacing = 1.25


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    first_line_indent_pt: int | None = None,
    bold: bool = False,
    size_pt: int = 12,
    before_pt: int = 0,
    after_pt: int = 0,
) -> None:
    """添加普通段落，并在一个地方统一控制字号、缩进和段前段后。"""

    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if first_line_indent_pt is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold)


def add_page_break(doc: Document) -> None:
    """显式插入分页，保证摘要、目录和正文各自独立成页。"""

    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_cover_page(doc: Document) -> None:
    """补一个更像正式毕业论文的封面页，用于提升模板前置部分的成稿感。"""

    add_paragraph(doc, UNIVERSITY_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=22, before_pt=70, after_pt=45)
    add_paragraph(doc, "毕业论文", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=20, before_pt=10, after_pt=70)
    add_paragraph(doc, TITLE_CN, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=18, before_pt=25, after_pt=90)
    add_paragraph(doc, COLLEGE_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, before_pt=10, after_pt=18)
    add_paragraph(doc, AUTHOR_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, before_pt=8, after_pt=18)
    add_paragraph(doc, STUDENT_ID, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, before_pt=8, after_pt=18)
    add_paragraph(doc, SUPERVISOR_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, before_pt=8, after_pt=18)
    add_paragraph(doc, SUBMIT_DATE, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, before_pt=70, after_pt=0)


def add_title_page(doc: Document) -> None:
    """生成中文题目、中文摘要和中文关键词页面。"""

    add_paragraph(doc, TITLE_CN, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=18, before_pt=30, after_pt=18)
    add_paragraph(doc, "摘  要", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=16, before_pt=12, after_pt=12)
    add_paragraph(doc, ABSTRACT_CN, first_line_indent_pt=24, size_pt=12)
    add_paragraph(doc, f"关键词：{'；'.join(KEYWORDS_CN)}", first_line_indent_pt=0, size_pt=12, before_pt=12)


def add_english_abstract_page(doc: Document) -> None:
    """生成英文题目、英文摘要和英文关键词页面。"""

    add_paragraph(doc, TITLE_EN, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=15, before_pt=24, after_pt=18)
    add_paragraph(doc, "Abstract", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=14, before_pt=12, after_pt=12)
    add_paragraph(doc, ABSTRACT_EN, first_line_indent_pt=24, size_pt=10)
    add_paragraph(doc, f"Keywords: {'; '.join(KEYWORDS_EN)}", first_line_indent_pt=0, size_pt=10, before_pt=8)


def add_statement_page(doc: Document) -> None:
    """按模板原意重建诚信声明页。"""

    add_paragraph(doc, STATEMENT_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=18, before_pt=60, after_pt=30)
    add_paragraph(doc, STATEMENT_BODY, first_line_indent_pt=24, size_pt=14)
    add_paragraph(doc, "", size_pt=12, before_pt=120)
    add_paragraph(doc, STATEMENT_SIGN, size_pt=14, before_pt=80)


def prepare_section_header_footer(section, *, show_header: bool, show_page_number: bool) -> None:
    """按分区控制页眉页脚，前置部分尽量简洁，正文保留论文头和页码。"""

    section.start_type = WD_SECTION_START.NEW_PAGE

    header = section.header
    header.is_linked_to_previous = False
    if not header.paragraphs:
        header.add_paragraph()
    for idx, paragraph in enumerate(header.paragraphs):
        paragraph.text = ""
        if idx == 0 and show_header:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(HEADER_TEXT)
            set_run_font(run, size_pt=10)

    footer = section.footer
    footer.is_linked_to_previous = False
    if not footer.paragraphs:
        footer.add_paragraph()
    footer.paragraphs[0].text = ""
    if show_page_number:
        add_page_number(footer.paragraphs[0])


def add_toc_page(doc: Document) -> None:
    """插入目录标题和目录占位符，后续由 Word COM 自动替换成正式目录域。"""

    add_paragraph(doc, "目  录", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=18, before_pt=30, after_pt=18)
    add_paragraph(doc, "[[TOC]]", first_line_indent_pt=0, size_pt=12)


def load_markdown_sections(markdown_text: str) -> tuple[str, str, list[str]]:
    """从主稿中提取正文区间与参考文献列表。"""

    body_start = markdown_text.find("## 1 绪论")
    ref_start = markdown_text.find("## 8 参考文献")
    if body_start == -1 or ref_start == -1:
        raise ValueError("未能在论文主稿中找到正文或参考文献标题。")

    body_text = markdown_text[body_start:ref_start].strip()
    ref_text = markdown_text[ref_start:].splitlines()
    references = [line.strip() for line in ref_text if re.match(r"^\[\d+\]", line.strip())]
    return markdown_text[:body_start], body_text, references


def strip_heading_numbering(text: str) -> str:
    """去掉标题文本中显式写出的层级编号，避免与模板自带标题编号重复。"""

    return re.sub(r"^\d+(?:\.\d+)*\s*", "", text).strip()


def clean_inline_markdown(text: str) -> str:
    """清理段落中的轻量 Markdown 标记，保留论文需要的正文文本。"""

    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = text.replace("\\", "")
    return text.strip()


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    """把 Markdown 表格行转成二维数组，供后续 DOCX 表格生成。"""

    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [clean_inline_markdown(item.strip()) for item in stripped.strip("|").split("|")]
        if cells and not all(re.fullmatch(r"[-:\s]+", cell) for cell in cells):
            rows.append(cells)
    return rows


def parse_markdown_blocks(body_text: str) -> list[dict]:
    """以当前项目主稿结构为前提，解析标题、正文、列表、表格和图片块。"""

    lines = body_text.splitlines()
    blocks: list[dict] = []
    paragraph_buffer: list[str] = []
    i = 0

    def flush_paragraph() -> None:
        if paragraph_buffer:
            text = clean_inline_markdown(" ".join(item.strip() for item in paragraph_buffer))
            if text:
                blocks.append({"type": "paragraph", "text": text})
            paragraph_buffer.clear()

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1)) - 1
            blocks.append({"type": "heading", "level": level, "text": clean_inline_markdown(heading_match.group(2))})
            i += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "table", "rows": parse_markdown_table(table_lines)})
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if image_match:
            flush_paragraph()
            image_path = image_match.group(2).lstrip("/")
            blocks.append(
                {
                    "type": "image",
                    "caption": clean_inline_markdown(image_match.group(1)),
                    "path": image_path,
                }
            )
            i += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match:
            flush_paragraph()
            items: list[str] = []
            while i < len(lines):
                inner = lines[i].strip()
                current_match = re.match(r"^\d+\.\s+(.*)$", inner)
                if not current_match:
                    break
                items.append(clean_inline_markdown(current_match.group(1)))
                i += 1
            blocks.append({"type": "list", "ordered": True, "items": items})
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(clean_inline_markdown(lines[i].strip()[2:]))
                i += 1
            blocks.append({"type": "list", "ordered": False, "items": items})
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()
    return blocks


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    """为表格单元格写入内容，并统一控制字号、对齐和行距。"""

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_cell_margin(table, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    """给表格单元格增加内边距，避免文字贴边导致阅读拥挤。"""

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)

    for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    """为表头单元格添加浅色底纹，提升表格层级感。"""

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table_with_caption(doc: Document, caption: str, rows: list[list[str]]) -> None:
    """在正文中插入表题和表格，并按论文场景做基础美化。"""

    add_paragraph(doc, caption, style="Caption" if "Caption" in doc.styles else None, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if "Table Grid" in [style.name for style in doc.styles]:
        table.style = "Table Grid"
    table.autofit = True
    set_table_cell_margin(table)

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), value, bold=(row_idx == 0))
            if row_idx == 0:
                shade_cell(table.cell(row_idx, col_idx), "D9E2F3")

    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    """插入图片并在下方添加图题。"""

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    width = Cm(15.8) if "montage" in image_path.name else Cm(14.6)
    run.add_picture(str(image_path), width=width)
    add_paragraph(doc, caption, style="Caption" if "Caption" in doc.styles else None, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11)


def add_list_block(doc: Document, items: Iterable[str], *, ordered: bool) -> None:
    """把 Markdown 列表转成论文中更稳妥的独立条目段落。"""

    for idx, item in enumerate(items, start=1):
        prefix = f"{idx}. " if ordered else "• "
        add_paragraph(doc, f"{prefix}{item}", first_line_indent_pt=0, size_pt=12)


def add_page_number(footer_paragraph) -> None:
    """在页脚中插入 PAGE 域，保证导出的 PDF 能看到页码。"""

    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run()
    set_run_font(run, size_pt=10)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def build_appendix_blocks(prompt_text: str) -> list[dict]:
    """把 LVLM 提示词文档转换成附录区块，便于直接复用现有整理成果。"""

    body_lines = []
    keep = False
    for line in prompt_text.splitlines():
        if line.startswith("## "):
            keep = True
        if keep:
            body_lines.append(line)
    return parse_markdown_blocks("\n".join(body_lines))


def count_document_characters(doc: Document) -> int:
    """统计文档中的非空白字符数，用于检查是否达到论文字数要求。"""

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return len(re.sub(r"\s+", "", text))


def build_document(args: argparse.Namespace) -> dict:
    """主流程：清空模板、副本写入、插图表格、附录与参考文献生成。"""

    global UNIVERSITY_NAME, COLLEGE_NAME, AUTHOR_NAME, STUDENT_ID, SUPERVISOR_NAME, SUBMIT_DATE, REFERENCE_HEADING

    UNIVERSITY_NAME = args.university_name
    COLLEGE_NAME = args.college_name
    AUTHOR_NAME = f"学生姓名：{args.author_name}"
    STUDENT_ID = f"学    号：{args.student_id}"
    SUPERVISOR_NAME = f"指导教师：{args.supervisor_name}"
    SUBMIT_DATE = args.submit_date
    REFERENCE_HEADING = args.reference_heading

    template_path = Path(args.template)
    source_md_path = Path(args.source_md)
    appendix_path = Path(args.appendix_md)
    output_path = Path(args.output)

    doc = Document(str(template_path))
    configure_styles(doc)
    remove_all_body_content(doc)
    prepare_section_header_footer(doc.sections[0], show_header=False, show_page_number=False)

    source_text = source_md_path.read_text(encoding="utf-8")
    _, body_text, references = load_markdown_sections(source_text)
    body_blocks = parse_markdown_blocks(body_text)
    appendix_blocks = build_appendix_blocks(appendix_path.read_text(encoding="utf-8"))

    add_cover_page(doc)
    add_page_break(doc)
    add_statement_page(doc)
    add_page_break(doc)
    add_title_page(doc)
    add_page_break(doc)
    add_english_abstract_page(doc)
    add_page_break(doc)
    add_toc_page(doc)
    add_page_break(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    prepare_section_header_footer(doc.sections[-1], show_header=True, show_page_number=True)

    pending_table_caption: str | None = None
    for block in body_blocks:
        block_type = block["type"]

        if block_type == "heading":
            heading_style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(block["level"], "Heading 3")
            add_paragraph(doc, strip_heading_numbering(block["text"]), style=heading_style, first_line_indent_pt=0, bold=True)
            pending_table_caption = None
            continue

        if block_type == "paragraph":
            text = block["text"]
            if re.fullmatch(r"表\s+\d+-\d+.*", text):
                pending_table_caption = text
            else:
                add_paragraph(doc, text, first_line_indent_pt=24, size_pt=12)
                pending_table_caption = None
            continue

        if block_type == "table":
            if not pending_table_caption:
                pending_table_caption = "实验结果表"
            add_table_with_caption(doc, pending_table_caption, block["rows"])
            pending_table_caption = None
            continue

        if block_type == "image":
            image_path = Path(block["path"])
            if not image_path.is_absolute():
                image_path = Path.cwd() / image_path
            add_figure(doc, image_path, block["caption"])
            pending_table_caption = None
            continue

        if block_type == "list":
            add_list_block(doc, block["items"], ordered=block["ordered"])
            pending_table_caption = None

    add_paragraph(doc, "致谢", style="Heading 1", first_line_indent_pt=0, bold=True)
    add_paragraph(doc, ACKNOWLEDGEMENTS, first_line_indent_pt=24, size_pt=12)

    add_paragraph(doc, "附录A LVLM案例分析提示词模板", style="Heading 1", first_line_indent_pt=0, bold=True)
    for block in appendix_blocks:
        if block["type"] == "heading":
            style = {1: "Heading 2", 2: "Heading 3", 3: "Heading 3"}.get(block["level"], "Heading 3")
            add_paragraph(doc, block["text"], style=style, first_line_indent_pt=0, bold=True)
        elif block["type"] == "paragraph":
            add_paragraph(doc, block["text"], first_line_indent_pt=24, size_pt=12)
        elif block["type"] == "list":
            add_list_block(doc, block["items"], ordered=block["ordered"])

    add_paragraph(doc, REFERENCE_HEADING, style="Heading 1", first_line_indent_pt=0, bold=True)
    for ref in references:
        style_name = "参考文献" if "参考文献" in doc.styles else None
        add_paragraph(doc, ref, style=style_name, first_line_indent_pt=0, size_pt=10 if style_name else 11)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    char_count = count_document_characters(doc)
    return {
        "output_docx": str(output_path.resolve()),
        "character_count_no_spaces": char_count,
        "reference_count": len(references),
        "body_block_count": len(body_blocks),
        "appendix_block_count": len(appendix_blocks),
        "university_name": UNIVERSITY_NAME,
        "college_name": COLLEGE_NAME,
        "author_name": AUTHOR_NAME,
        "student_id": STUDENT_ID,
        "supervisor_name": SUPERVISOR_NAME,
        "submit_date": SUBMIT_DATE,
    }


def parse_args() -> argparse.Namespace:
    """解析命令行参数，便于后续继续复用这个生成流程。"""

    parser = argparse.ArgumentParser(description="基于模板副本生成项目论文 DOCX。")
    parser.add_argument("--template", required=True, help="作为样式和页面设置基础的模板 DOCX 路径。")
    parser.add_argument("--source_md", default="docs/full_paper_draft.md", help="论文主稿 Markdown 路径。")
    parser.add_argument("--appendix_md", default="docs/lvlm_prompt_templates.md", help="附录提示词 Markdown 路径。")
    parser.add_argument("--output", required=True, help="输出 DOCX 路径。")
    parser.add_argument("--report", required=True, help="输出生成报告 JSON 路径。")
    parser.add_argument("--university_name", default="本科毕业论文", help="封面顶部标题。")
    parser.add_argument("--college_name", default="人工智能与计算机相关专业方向", help="学院或专业方向。")
    parser.add_argument("--author_name", default="__________", help="学生姓名。")
    parser.add_argument("--student_id", default="__________", help="学生学号。")
    parser.add_argument("--supervisor_name", default="__________", help="指导教师姓名。")
    parser.add_argument("--submit_date", default="完成时间：2026 年 4 月", help="封面完成时间。")
    parser.add_argument("--reference_heading", default="参考文献", help="参考文献章节标题。")
    return parser.parse_args()


def main() -> None:
    """程序入口：执行生成并把统计信息写到 JSON 报告。"""

    args = parse_args()
    report = build_document(args)

    report_path = Path(args.report)
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
