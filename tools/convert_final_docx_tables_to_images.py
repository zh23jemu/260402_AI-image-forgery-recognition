"""把论文中的 Word 原生表格转换为稳定的 PNG 表格图片。

artifact-tool 对当前模板中的 Word 表格渲染不稳定，可能把表格压成竖排。
为了保证最终版在渲染检查中不出现表格跑版，本脚本执行三步：
1. 从 Word 原生表格读取真实内容；
2. 按论文风格生成高分辨率 PNG 表格图片；
3. 在原表格位置插入图片，并移除原生表格对象，保留原来的表题编号。

脚本只修改 docx 内部的表格呈现方式，不改正文论述和已有图注。
"""

from __future__ import annotations

import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "AI伪造图像识别论文_终稿.docx"
FIGURES_DIR = ROOT / "figures"


TABLE_IMAGE_CONFIGS = [
    {
        "filename": "table_5_1_fsd_results.png",
        "title": "表 5-1 FSD 基线实验结果",
        "width": 1500,
        "col_weights": [1.9, 1.2, 1.2, 1.2],
    },
    {
        "filename": "table_5_2_stay_positive_results.png",
        "title": "表 5-2 Stay-Positive 基线实验结果",
        "width": 1500,
        "col_weights": [1.7, 1.4, 1.2, 1.2],
    },
    {
        "filename": "table_5_3_fsd_training_exploration.png",
        "title": "表 5-3 FSD 训练探索结果",
        "width": 1800,
        "col_weights": [1.8, 1.35, 1.65, 1.25, 1.25, 1.25],
    },
    {
        "filename": "table_5_4_joint_training_results.png",
        "title": "表 5-4 两阶段联合训练结果对比",
        "width": 1900,
        "col_weights": [2.0, 1.2, 1.2, 1.6, 0.9, 2.8],
    },
]


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    """优先使用 Windows 中文字体，确保中文表头和正文不乱码。"""

    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


def split_text_for_width(text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    """按照像素宽度对中英文混排文本换行，避免长单元格挤出边界。"""

    text = text.strip()
    if not text:
        return [""]

    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if font.getlength(candidate) <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_centered_multiline(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    lines: list[str],
    font: ImageFont.FreeTypeFont,
    fill: tuple[int, int, int],
) -> None:
    """在单元格中垂直、水平居中绘制多行文字。"""

    left, top, right, bottom = box
    line_height = int(font.size * 1.25)
    total_height = line_height * len(lines)
    y = top + max(0, (bottom - top - total_height) // 2)
    for line in lines:
        text_width = font.getlength(line)
        x = left + max(0, int((right - left - text_width) / 2))
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def render_table_png(rows: list[list[str]], config: dict) -> Path:
    """把表格数据渲染成论文可用的 PNG 图片。"""

    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    output_path = FIGURES_DIR / config["filename"]

    width = int(config["width"])
    margin_x = 46
    margin_y = 34
    title_height = 54
    footer = 22
    header_font = load_font(28, bold=True)
    body_font = load_font(25)
    title_font = load_font(30, bold=True)

    usable_width = width - margin_x * 2
    weights = config["col_weights"]
    weight_sum = sum(weights)
    col_widths = [int(usable_width * weight / weight_sum) for weight in weights]
    col_widths[-1] += usable_width - sum(col_widths)

    wrapped_rows: list[list[list[str]]] = []
    row_heights: list[int] = []
    for row_index, row in enumerate(rows):
        font = header_font if row_index == 0 else body_font
        wrapped_cells = []
        max_lines = 1
        for col_index, value in enumerate(row):
            max_text_width = col_widths[col_index] - 26
            lines = split_text_for_width(value, font, max_text_width)
            wrapped_cells.append(lines)
            max_lines = max(max_lines, len(lines))
        wrapped_rows.append(wrapped_cells)
        base_height = 70 if row_index == 0 else 66
        row_heights.append(max(base_height, int(font.size * 1.28) * max_lines + 24))

    height = margin_y * 2 + title_height + sum(row_heights) + footer
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    # 标题直接画进图片顶部，后续 Word 中仍保留原表题编号，双重保障读者识别表格内容。
    title = config["title"]
    title_width = title_font.getlength(title)
    draw.text(((width - title_width) / 2, margin_y - 4), title, font=title_font, fill=(25, 40, 55))

    x_edges = [margin_x]
    for col_width in col_widths:
        x_edges.append(x_edges[-1] + col_width)

    y = margin_y + title_height
    border = (60, 80, 100)
    header_bg = (218, 234, 247)
    alt_bg = (247, 250, 253)

    for row_index, wrapped_cells in enumerate(wrapped_rows):
        row_height = row_heights[row_index]
        bg = header_bg if row_index == 0 else (alt_bg if row_index % 2 == 0 else (255, 255, 255))
        draw.rectangle([margin_x, y, width - margin_x, y + row_height], fill=bg)
        for col_index, lines in enumerate(wrapped_cells):
            left = x_edges[col_index]
            right = x_edges[col_index + 1]
            draw.rectangle([left, y, right, y + row_height], outline=border, width=2)
            font = header_font if row_index == 0 else body_font
            fill = (20, 35, 50) if row_index == 0 else (30, 30, 30)
            draw_centered_multiline(draw, (left + 8, y + 5, right - 8, y + row_height - 5), lines, font, fill)
        y += row_height

    # 外边框稍微加重，避免缩放进 Word 后表格边界过淡。
    draw.rectangle([margin_x, margin_y + title_height, width - margin_x, y], outline=border, width=3)
    image.save(output_path, quality=95)
    return output_path


def insert_image_before_table(table, image_path: Path, width_cm: float) -> None:
    """在原表格位置前插入图片段落，然后删除原生表格。"""

    new_p = OxmlElement("w:p")
    table._tbl.addprevious(new_p)
    paragraph = Paragraph(new_p, table._parent)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    table_element = table._element
    table_element.getparent().remove(table_element)


def main() -> None:
    """执行表格图片化转换。"""

    doc = Document(DOCX_PATH)
    if len(doc.tables) == 0:
        print("tables=0; no conversion needed")
        return

    table_rows = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        table_rows.append(rows)

    image_paths = []
    for rows, config in zip(table_rows, TABLE_IMAGE_CONFIGS):
        image_paths.append(render_table_png(rows, config))

    # 需要先复制列表，因为删除表格后 doc.tables 会动态变化。
    for table, image_path, config in zip(list(doc.tables), image_paths, TABLE_IMAGE_CONFIGS):
        width_cm = 15.0 if config["width"] <= 1500 else 15.6
        insert_image_before_table(table, image_path, width_cm)

    doc.save(DOCX_PATH)
    print(f"converted_tables={len(image_paths)}")
    for image_path in image_paths:
        print(image_path)


if __name__ == "__main__":
    main()
